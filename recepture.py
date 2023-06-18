import logging
import os
import shutil
import traceback
from collections import Counter
from decimal import Decimal
from functools import reduce
from textwrap import wrap

import docx
from win32com import client

import numpy as np
from PyQt6 import QtCore, QtGui, QtWidgets
from PyQt6.QtCore import Qt, QSize
from PyQt6.QtWidgets import QCompleter, QFileDialog
from sqlitedict import SqliteDict

from common.secrets import Secrets
from common.ui_elements import HoverableButton, MenuButton, ColorButton, CustomMenu, CustomRadioBtn, generate_color, \
    CustomListItem, generate_font, MplCanvas, delete_chield, create_w_lo, normalize_number, get_numeric_validator, \
    insert_w_lo, get_h_spacer, get_v_spacer, change_position_window, ChoiceColor, set_window_icon
from component_card import CustomEntry, CustomCombobox
from database import DB
from typing import List, Tuple
import xml.etree.ElementTree as ET
import requests
from newReactives import InfoWindow, DarkBtn_Ui
from common.settings import get_suhoi_type, update_config_param
from skimage import color as color_kit



class ReceptureWindow(QtWidgets.QWidget):

    def __init__(self, project_name: str, iter_name: str, name: str, project_window=None, is_new=False):
        super(ReceptureWindow, self).__init__()
        self.project_window = project_window
        self.project = project_name
        self.iter = iter_name
        self.name = name
        self.db = DB()
        self.setWindowTitle(f"{self.project} - {self.iter} - {self.name}")
        self.recepture_data = ReceptureDataModel(project_name, iter_name, name)
        self.recepture_data.load_data()
        self.settings_window = None
        self.additive_window = None
        self.hardener_window = None
        self.recount_on_maslo_window = None
        self.count_recepture_constant = None
        self.count_recepture_combo = None
        self.philum_window = None
        self.choice_color_window = None
        self.save_as_window = None
        self.list_comp_row_obj = []
        self.list_comp_2_row_obj = []
        self.list_experiment_obj = []

        set_window_icon(self)
        self.verticalLayout_3 = QtWidgets.QVBoxLayout(self)
        self.verticalLayout_3.setContentsMargins(0, 0, 0, 0)
        self.verticalLayout_3.setSpacing(0)
        toolbar = self.add_toolbar(self)
        self.verticalLayout_3.addWidget(toolbar)

        self.tabWidget = QtWidgets.QTabWidget(parent=self)
        self.tabWidget.setObjectName("tabWidget")

        self.recepture_tab = QtWidgets.QWidget()
        self.verticalLayout_2 = QtWidgets.QVBoxLayout(self.recepture_tab)
        self.verticalLayout_2.setContentsMargins(9,0,9,9)
        self.verticalLayout_2.setSpacing(0)


        self.widget = QtWidgets.QWidget(parent=self.recepture_tab)
        self.widget.setObjectName("widget")
        self.horizontalLayout_6 = QtWidgets.QHBoxLayout(self.widget)
        self.horizontalLayout_6.setObjectName("horizontalLayout_6")

        self.left_side = QtWidgets.QWidget(parent=self.widget)
        self.left_side.setMaximumSize(QtCore.QSize(390, 16777215))
        self.horizontalLayout_6.addWidget(self.left_side)
        self.left_vertical_lo = QtWidgets.QVBoxLayout(self.left_side)
        self.left_vertical_lo.setContentsMargins(0,0,0,0)
        self.left_vertical_lo.setSpacing(0)

        self.scrollArea = QtWidgets.QScrollArea(parent=self.widget)
        self.scrollArea.setMinimumSize(QtCore.QSize(390, 0))
        self.scrollArea.setMaximumSize(QtCore.QSize(390, 16777215))
        self.scrollArea.setWidgetResizable(True)
        self.scrollArea.setObjectName("scrollArea")
        self.widget.setStyleSheet("""
                        QWidget#recepture{
                                  background: #f9f9f9;
                                  border: 0px solid black;
                                  }
                        QScrollArea#scrollArea{
                           background: #f9f9f9;
                           border: 0px solid #bbb;
                           }          
                                  """)
        self.recepture = QtWidgets.QWidget()
        self.recepture.setGeometry(QtCore.QRect(0, 0, 403, 485))
        self.recepture.setObjectName("recepture")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.recepture)
        self.verticalLayout.setContentsMargins(0,0,0,0)
        self.verticalLayout.setSpacing(2)

        self.component_one = Ui_Component(self.recepture, self.recepture_data)
        # self.component_one.setGeometry(QtCore.QRect(0, 0, 403, 485))
        # self.component_one.setObjectName("recepture")
        # self.component_one_l = QtWidgets.QVBoxLayout(self.recepture)
        # self.component_one_l.setContentsMargins(0, 0, 0, 0)
        # self.component_one_l.setSpacing(2)

        self.buttons = QtWidgets.QWidget(parent=self.recepture)
        self.buttons.setObjectName("buttons")
        self.horizontalLayout_7 = QtWidgets.QHBoxLayout(self.buttons)
        self.horizontalLayout_7.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout_7.setObjectName("horizontalLayout_7")
        self.btn_2k = HoverableButton(self.buttons, "2k", (20, 20))
        self.btn_2k.clicked.connect(self.show_2k)
        self.horizontalLayout_7.addWidget(self.btn_2k, alignment=QtCore.Qt.AlignmentFlag.AlignRight)
        self.plus = HoverableButton(self.buttons, "plus_2", (20, 20))
        self.plus.clicked.connect(lambda x: self.add_row("one"))
        self.horizontalLayout_7.addWidget(self.plus, alignment=QtCore.Qt.AlignmentFlag.AlignLeft)


        self.verticalLayout.addWidget(self.component_one)
        self.verticalLayout.addWidget(self.buttons)

        self.component_two = Ui_Component(self.recepture, self.recepture_data)
        # self.component_two.setGeometry(QtCore.QRect(0, 0, 403, 485))
        # self.component_two.setObjectName("recepture")
        # self.component_two_l = QtWidgets.QVBoxLayout(self.component_two)
        # self.component_two_l.setContentsMargins(0, 0, 0, 0)
        # self.component_two_l.setSpacing(2)

        self.plus_2 = MenuButton(self.recepture, "plus_2", (20, 20))
        self.plus_2.clicked.connect(lambda x: self.add_row("two"))
        self.plus_2.hide()
        self.verticalLayout.addWidget(self.component_two)
        self.verticalLayout.addWidget(self.plus_2, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)
        spacerItem1 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                            QtWidgets.QSizePolicy.Policy.Expanding)
        self.verticalLayout.addItem(spacerItem1)

        for (name, value), comment in zip(self.recepture_data.component_list, self.recepture_data.list_comments):
             self.add_row("one", name=name, value=value, comment=comment)
        for (name, value), comment in zip(self.recepture_data.component_list_2, self.recepture_data.list_comments_2):
            self.add_row("two", name=name, value=value, comment=comment)

        if not self.recepture_data.flag_2k:
            self.component_two.hide()

        self.scrollArea.setWidget(self.recepture)
        self.left_vertical_lo.addWidget(self.scrollArea)

        self.right_side = QtWidgets.QWidget(parent=self.widget)
        self.verticalLayout_6 = QtWidgets.QVBoxLayout(self.right_side)
        self.verticalLayout_6.setSpacing(3)
        self.verticalLayout_6.setContentsMargins(0,0,0,0)
        self.right_side.setMaximumSize(350, 999999)

        w, lo = create_w_lo(self.right_side, self.verticalLayout_6)
        self.count_params_l = QtWidgets.QLabel(parent=w)
        self.count_params_l.setText("Расчетные параметры")
        self.count_params_l.setFont(generate_font(12))
        lo.addWidget(self.count_params_l)
        self.setting_count = HoverableButton(w, "settings", (16,16))
        self.setting_count.clicked.connect(lambda: self.open_r_settings())
        lo.addWidget(self.setting_count, alignment=QtCore.Qt.AlignmentFlag.AlignRight)
        lo.setContentsMargins(0,0,0,9)


        self.price_l = QtWidgets.QLabel(parent=self.right_side)
        self.price_l.setText("Стоимость:")
        self.verticalLayout_6.addWidget(self.price_l)

        self.density_l = QtWidgets.QLabel(parent=self.right_side)
        self.density_l.setText("Плотность:")
        self.verticalLayout_6.addWidget(self.density_l)

        w, lo = create_w_lo(self.right_side, self.verticalLayout_6)
        self.suhoi_l = QtWidgets.QLabel(parent=w)
        self.suhoi_l.setText("Масс.д.н.в:")
        lo.addWidget(self.suhoi_l)

        self.volume_suhoi_l = QtWidgets.QLabel(parent=w)
        self.volume_suhoi_l.setText("Объем.д.н.в:")
        lo.addWidget(self.volume_suhoi_l)


        self.oil_l = QtWidgets.QLabel(parent=self.right_side)
        self.oil_l.setText("Маслоемкость 1-го рода:")
        self.verticalLayout_6.addWidget(self.oil_l)

        self.philum_l = QtWidgets.QLabel(parent=self.right_side)
        self.philum_l.setText("Филум пигментов:")
        self.verticalLayout_6.addWidget(self.philum_l)

        self.degree_pigm_l = QtWidgets.QLabel(parent=self.right_side)
        self.degree_pigm_l.setText("Степень пигментирования:")
        self.verticalLayout_6.addWidget(self.degree_pigm_l)

        self.const_pigm_l = QtWidgets.QLabel(parent=self.right_side)
        self.const_pigm_l.setText("Константа наполнения:")
        self.verticalLayout_6.addWidget(self.const_pigm_l)

        w, lo = create_w_lo(self.right_side, self.verticalLayout_6)
        self.okp_l = QtWidgets.QLabel(parent=w)
        self.okp_l.setText("ОКП:")
        lo.addWidget(self.okp_l)

        self.kokp_l = QtWidgets.QLabel(parent=w)
        self.kokp_l.setText("КОКП:")
        lo.addWidget(self.kokp_l)

        self.okp_kokp_l = QtWidgets.QLabel(parent=w)
        self.okp_kokp_l.setText("ОКП/КОКП:")
        lo.addWidget(self.okp_kokp_l)

        self.hiding_pigm_l = QtWidgets.QLabel(parent=self.right_side)
        self.hiding_pigm_l.setText("Укрывистость пигментов:")
        self.verticalLayout_6.addWidget(self.hiding_pigm_l)

        self.hiding_wet_l = QtWidgets.QLabel(parent=self.right_side)
        self.hiding_wet_l.setText("Укрывистость мокрой пленки:")
        self.verticalLayout_6.addWidget(self.hiding_wet_l)

        self.hiding_dry_l = QtWidgets.QLabel(parent=self.right_side)
        self.hiding_dry_l.setText("Укрывистость сухой пленки:")
        self.verticalLayout_6.addWidget(self.hiding_dry_l)

        w, lo = create_w_lo(self.right_side, self.verticalLayout_6)
        self.count_btn = DarkBtn_Ui(w, "calc")
        self.count_btn.clicked.connect(self.count_all)
        lo.addWidget(self.count_btn)
        lo.setContentsMargins(0,9,0,9)


        l = QtWidgets.QLabel(parent=self.right_side)
        l.setText("Дополнительные функции")
        l.setFont(generate_font(11))
        self.verticalLayout_6.addWidget(l,alignment=QtCore.Qt.AlignmentFlag.AlignLeft)

        w, lo = create_w_lo(self.right_side, self.verticalLayout_6)
        self.count_components_btn = ColorButton(w,  "blue")
        self.count_components_btn.setText("Расчет компонентов")
        menu = CustomMenu(self)
        menu.addAction('Расчет функц. добавок', lambda: self.open_count_additives())
        menu.addAction('Расчет отвердителя', lambda: self.open_count_hardeners())
        menu.addAction('Заменить по маслоемкости', lambda: self.open_recount_on_maslo())
        self.count_components_btn.setMenu(menu)
        lo.addWidget(self.count_components_btn)

        w, lo = create_w_lo(self.right_side, self.verticalLayout_6)
        self.count_new_recepture = ColorButton(w,  "blue")
        self.count_new_recepture.setText("Расчет рецептур")
        menu = CustomMenu(self)
        menu.addAction('По константе наполнения', lambda: self.open_count_recepture_const())
        menu.addAction('Комбинированный расчет', lambda: self.open_count_recepture_comb())
        self.count_new_recepture.setMenu(menu)
        lo.addWidget(self.count_new_recepture)

        w, lo = create_w_lo(self.right_side, self.verticalLayout_6)
        self.others = ColorButton(w,  "blue")
        self.others.setText("Разное")
        menu = CustomMenu(self)
        menu.addAction('Филумы пигментов', lambda: self.open_philums())
        self.others.setMenu(menu)
        lo.addWidget(self.others)


        spacerItem2 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                            QtWidgets.QSizePolicy.Policy.Expanding)
        self.verticalLayout_6.addItem(spacerItem2)
        self.horizontalLayout_6.addWidget(self.right_side)
        self.verticalLayout_2.addWidget(self.widget)

        self.all_amount_w = QtWidgets.QWidget(parent=self.left_side)
        space = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Expanding,
                                            QtWidgets.QSizePolicy.Policy.Minimum)

        self.horizontalLayout_8 = QtWidgets.QHBoxLayout(self.all_amount_w)
        self.horizontalLayout_8.setContentsMargins(0,9,30,9)
        self.horizontalLayout_8.setSpacing(9)
        self.horizontalLayout_8.addItem(space)
        self.amount_all_l = QtWidgets.QLabel(parent=self.all_amount_w)
        self.amount_all_l.setText("Итого:")
        self.horizontalLayout_8.addWidget(self.amount_all_l)
        self.amount_all_value = QtWidgets.QLabel(parent=self.all_amount_w)
        self.horizontalLayout_8.addWidget(self.amount_all_value)
        self.recount_btn = HoverableButton(self.all_amount_w, "menu", (16,16))
        menu = CustomMenu(self)

        menu.addAction('Списать компоненты', lambda : self.subtract_warehouse())
        menu.addAction('Пересчитать массу', lambda : self.recount_mass())
        menu.addAction('Довести растворителем', lambda :self.add_solvent())

        self.recount_btn.setMenu(menu)

        self.horizontalLayout_8.addWidget(self.recount_btn)
        self.left_vertical_lo.addWidget(self.all_amount_w)
        self.tabWidget.addTab(self.recepture_tab, "Рецептура")


        self.experimental_tab = QtWidgets.QWidget()
        self.verticalLayout_4 = QtWidgets.QVBoxLayout(self.experimental_tab)
        self.verticalLayout_4.setContentsMargins(9, 0, 9, 9)
        self.exp_body_w = QtWidgets.QWidget(parent=self.experimental_tab)
        self.exp_body_w.setObjectName("exp_body_w")
        self.horizontalLayout_2 = QtWidgets.QHBoxLayout(self.exp_body_w)
        self.horizontalLayout_2.setObjectName("horizontalLayout_2")
        self.exp_params_w = QtWidgets.QWidget(parent=self.exp_body_w)
        self.exp_params_w.setObjectName("exp_params_w")
        self.gridLayout_2 = QtWidgets.QGridLayout(self.exp_params_w)
        self.gridLayout_2.setObjectName("gridLayout_2")

        l = QtWidgets.QLabel(parent=self.exp_params_w)
        l.setText("Экспериментальные значения")
        l.setFont(generate_font(12))
        self.gridLayout_2.addWidget(l, 0, 0, 1, 2)

        self.scrollArea_2 = QtWidgets.QScrollArea(parent=self.exp_params_w)
        self.scrollArea_2.setMinimumSize(QtCore.QSize(475, 350))
        self.scrollArea_2.setMaximumSize(QtCore.QSize(700, 16777215))
        self.scrollArea_2.setWidgetResizable(True)
        self.scrollArea_2.setObjectName("scrollArea")
        self.exp_params_w.setStyleSheet("""
                                QWidget#recepture{
                                          background: #f9f9f9;
                                          border: 0px solid black;
                                          }
                                QScrollArea#scrollArea{
                                   background: #f9f9f9;
                                   border: 0px solid #bbb;
                                   }          
                                          """)
        self.exp_s_area = QtWidgets.QWidget()
        self.exp_s_area.setGeometry(QtCore.QRect(0, 0, 500, 700))
        self.exp_s_area.setObjectName("recepture")
        self.gridLayout_3 = QtWidgets.QGridLayout(self.exp_s_area)
        self.gridLayout_3.setVerticalSpacing(3)
        l = QtWidgets.QLabel(parent=self.exp_s_area)
        l.setText("Название")
        l.setFont(generate_font(10))
        self.gridLayout_3.addWidget(l, 1, 0, 1, 1)

        l = QtWidgets.QLabel(parent=self.exp_s_area)
        l.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        l.setText("Требуемое \n"  "значение")
        l.setFont(generate_font(10))
        self.gridLayout_3.addWidget(l, 1, 1, 1, 1)

        l = QtWidgets.QLabel(parent=self.exp_s_area)
        l.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        l.setText("Полученное \nзначение")
        l.setFont(generate_font(10))
        self.gridLayout_3.addWidget(l, 1, 2, 1, 1)

        l = QtWidgets.QLabel(parent=self.exp_s_area)
        l.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        l.setText("Успех")
        l.setFont(generate_font(10))
        self.gridLayout_3.addWidget(l, 1, 3, 1, 3)

        for row in self.recepture_data.experiment_list:
            self.add_exp_row(*row)

        spacerItem3 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                            QtWidgets.QSizePolicy.Policy.Expanding)
        self.gridLayout_3.addItem(spacerItem3, 100, 0, 1, 1)

        self.scrollArea_2.setWidget(self.exp_s_area)
        self.gridLayout_2.addWidget(self.scrollArea_2, 2, 0, 1, 4)

        spacerItem3 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                            QtWidgets.QSizePolicy.Policy.Expanding)
        self.gridLayout_2.addItem(spacerItem3, 100, 0, 1, 1)
        self.horizontalLayout_2.addWidget(self.exp_params_w)

        self.color_w = QtWidgets.QWidget(parent=self.exp_body_w)
        self.gridLayout_4 = QtWidgets.QGridLayout(self.color_w)

        self.lable_name = QtWidgets.QLabel(parent=self.color_w)

        self.lable_name.setFont(generate_font(12))
        self.lable_name.setText("Цвет")
        self.gridLayout_4.addWidget(self.lable_name, 0, 0, 1, 1)

        self.lable_color1 = QtWidgets.QLabel(parent=self.color_w)
        self.lable_color1.setText("Требуемый цвет")
        self.gridLayout_4.addWidget(self.lable_color1, 1, 0, 1, 1)

        self.color1 = QtWidgets.QLabel(parent=self.color_w)
        image = QtGui.QPixmap(generate_color(self.recepture_data.project_color))
        self.color1.setPixmap(image)
        self.color1.setMaximumSize(QSize(82, 80))
        self.color1.setStyleSheet("""
        QLabel{
        border: 1px solid #ddd;
        }
        """)
        self.gridLayout_4.addWidget(self.color1, 2, 0, 1, 1, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)

        self.lable_color2 = QtWidgets.QLabel(parent=self.color_w)
        self.lable_color2.setText("Полученный цвет")
        self.gridLayout_4.addWidget(self.lable_color2, 1, 1, 1, 1, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)

        self.color2 = QtWidgets.QLabel(parent=self.color_w)
        self.color2.setMaximumSize(QSize(82, 80))
        self.color2.setStyleSheet("""
        QLabel{
        border: 1px solid #ddd;
        }
        """)
        self.gridLayout_4.addWidget(self.color2, 2, 1, 1, 1, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)

        self.select_color_btn = ColorButton(self.color_w, "blue")
        self.select_color_btn.setText("Выбрать цвет")
        self.select_color_btn.clicked.connect(lambda: self.open_choice_color())
        self.gridLayout_4.addWidget(self.select_color_btn, 5, 1, 1, 1)

        self.delta_e_color_l = QtWidgets.QLabel(self.color_w)
        self.delta_e_color_l.setText("ΔЕ:")
        self.gridLayout_4.addWidget(self.delta_e_color_l, 6, 1, 1, 1, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)
        self.set_selected_color(self.recepture_data.recepture_color)

        spacerItem4 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                            QtWidgets.QSizePolicy.Policy.Expanding)
        self.gridLayout_4.addItem(spacerItem4, 7, 0, 1, 1)


        self.horizontalLayout_2.addWidget(self.color_w)
        spacerItem5 = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Policy.Expanding,
                                            QtWidgets.QSizePolicy.Policy.Minimum)
        self.horizontalLayout_2.addItem(spacerItem5)
        self.verticalLayout_4.addWidget(self.exp_body_w)
        self.tabWidget.addTab(self.experimental_tab, "Эксперимент")

        self.description_tab = QtWidgets.QWidget()
        self.verticalLayout_5 = QtWidgets.QVBoxLayout(self.description_tab)
        self.verticalLayout_5.setContentsMargins(9, 0, 9, 9)
        self.description = QtWidgets.QTextEdit(parent=self.description_tab)
        self.description.setText(self.recepture_data.notes)
        self.verticalLayout_5.addWidget(self.description)
        self.tabWidget.addTab(self.description_tab, "Заметки")
        self.verticalLayout_3.addWidget(self.tabWidget)

        self.tabWidget.setCurrentIndex(0)
        if not is_new:
            self.count_mass()

    def closeEvent(self, event):
        Ui_Component.list_obj.remove(self.component_one)
        Ui_Component.list_obj.remove(self.component_two)

    def add_toolbar(self, parent: QtWidgets) -> QtWidgets.QFrame:
        toolbar = QtWidgets.QFrame(parent=parent)
        toolbar.setObjectName("toolbarBg")
        toolbar.setStyleSheet("""
        QFrame#toolbarBg{
        }
        
        """)
        toolbar.setContentsMargins(9,0,9,0)

        horizontalLayout_3 = QtWidgets.QHBoxLayout(toolbar)
        horizontalLayout_3.setSpacing(4)
        self.name_recepture = QtWidgets.QLabel(parent=toolbar)
        self.name_recepture.setContentsMargins(5,0,25,0)
        self.name_recepture.setText(self.name)
        horizontalLayout_3.addWidget(self.name_recepture)
        self.name_recepture.setFont(generate_font(18))
        save_btn = HoverableButton(toolbar, "save", (20,20))
        save_btn.clicked.connect(lambda: self.save())
        horizontalLayout_3.addWidget(save_btn)
        save_as_btn = HoverableButton(toolbar, "save_as", (20,20))
        save_as_btn.clicked.connect(lambda: self.save(save_as=True))
        horizontalLayout_3.addWidget(save_as_btn)
        word_btn = HoverableButton(toolbar, "word", (20,20))
        word_btn.clicked.connect(lambda: self.save_word())
        horizontalLayout_3.addWidget(word_btn)

        spacerItem = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Policy.Expanding,
                                           QtWidgets.QSizePolicy.Policy.Minimum)
        horizontalLayout_3.addItem(spacerItem)
        delete_btn = HoverableButton(toolbar, "del_rec", (20,20))
        delete_btn.clicked.connect(lambda: self.delete())
        horizontalLayout_3.addWidget(delete_btn)

        return toolbar

    def show_2k(self, event):
        check = self.recepture_data.flag_2k
        if check:
            self.component_two.hide()
            self.plus_2.hide()
            self.recepture_data.flag_2k = False
        else:
            self.component_two.show()
            self.plus_2.show()
            self.recepture_data.flag_2k = True
        self.btn_2k.set_pressed()
        self.count_mass()

    def add_row(self, _type, name="", value="", comment="$None"):
        if _type == "one":
            self.component_one.add_row(name=name, value=value, callback_mass=self.count_mass, comment=comment)
        else:
            self.component_two.add_row(name=name, value=value, callback_mass=self.count_mass, comment=comment)

        # parent = self.component_one if _type == "one" else self.component_two
        # list_obj = self.list_comp_row_obj if _type == "one" else self.list_comp_2_row_obj
        # loyout = self.component_one_l if _type == "one" else self.component_two_l
        # add_widget = self.buttons if _type == "one" else self.plus_2
        # loyout.removeWidget(add_widget)
        # _index = len(list_obj)
        # row = ComponentRow(parent, _index, name=name, amount=value, list_obj=list_obj, callback_mass=self.count_mass)
        # loyout.addWidget(row)
        # loyout.addWidget(add_widget)
        # list_obj.append(row)
        #
        # self.reset_row_number(_type)

    def add_exp_row(self, name, needed, value, state):
        self.list_experiment_obj.append(ExperimentRow(self.exp_s_area, self.gridLayout_3, name, needed, value, state))

    def reset_row_number(self, _type: str):
        list_obj: List[ComponentRow]
        list_obj = self.list_comp_row_obj if _type == "one" else self.list_comp_2_row_obj
        number = 1
        for obj in list_obj:
            if obj is not None:
                obj.set_number(number)
                number += 1

    def collect_rows_data(self):
        list_1 = list(filter(lambda x: x!= None, self.component_one.get_list_obj()))
        list_2 = list(filter(lambda x: x!= None, self.component_two.get_list_obj()))

        list_comp_1 = [i.get_data() for i in list_1]
        list_comp_category_1 = [i.get_category() for i in list_1]
        list_comp_2 = [i.get_data() for i in list_2]
        list_comp_category_2 = [i.get_category() for i in list_2]

        self.recepture_data.component_list = list_comp_1
        self.recepture_data.component_list_2 = list_comp_2
        self.recepture_data.category_list = list_comp_category_1
        self.recepture_data.category_list_2 = list_comp_category_2

    def collect_exp_data(self):
        list_data = []
        for row in self.list_experiment_obj:
            list_data.append(row.get())

        self.recepture_data.experiment_list = list_data

    def collect_note(self):
        self.recepture_data.notes = self.description.toPlainText()

    def save(self, save_as=None):
        self.collect_rows_data()
        self.collect_exp_data()
        self.collect_note()
        if save_as is not None:
            if self.save_as_window is None:
                self.save_as_window = SaveAsWindow(self, self.project, self.iter, self.save_as)
                self.save_as_window.show()

        else:
            self.recepture_data.save()
            self.project_window.select_project(self.project)

    def save_as(self, iter, new_name):
        self.name_recepture.setText(new_name)
        self.recepture_data.save(save_as=(iter, new_name))
        self.project_window.select_project(self.project)

    def save_word(self):
        self.collect_rows_data()
        self.collect_exp_data()
        self.collect_note()
        WordExport(self.recepture_data)

    def delete(self):
        if InfoWindow(f"Вы уверены, что хотите удалить \nрецептуру: {self.name}").exec():
            self.recepture_data.delete()
            self.closeEvent(None)
            self.project_window.select_project(self.project)
            self.destroy()

    def count_mass(self, collected=False):
        if not collected:
            self.collect_rows_data()
        mass = self.recepture_data.count_mass(all=True)
        mass = normalize_number(mass)
        self.amount_all_value.setText(mass)

    def count_all(self):
        self.collect_rows_data()
        self.count_mass(collected=True)
        self.recepture_data.all_count()
        data = self.recepture_data
        list_update_lable_value = (
            (self.price_l, data.price, "руб/кг"),
            (self.oil_l, data.oil, "г/100 г"),
            (self.suhoi_l, data.suhoi, "%"),
            (self.volume_suhoi_l, data.volume_suhoi, "%"),
            (self.okp_l, data.okp, "%"),
            (self.kokp_l, data.kokp, "%"),
            (self.okp_kokp_l, data.okp_kokp, "%"),
            (self.hiding_pigm_l, data.hiding_pigm, "г/м²"),
            (self.hiding_wet_l, data.hiding_wet, "г/м²"),
            (self.hiding_dry_l, data.hiding_dry, "г/м²"),
            (self.philum_l, data.philum, ""),
            (self.density_l, data.get_density(), "г/см³"),
            (self.degree_pigm_l, data.degree_pigm, ""),
            (self.const_pigm_l, data.const_pigm, ""),
        )
        for lable, value, size in list_update_lable_value:
            self.update_lable_param(lable, value, size)

    def update_lable_param(self, lable: QtWidgets.QLabel, new_value: Decimal, size: str):
        # print(f"lable { lable.text()} new_value {new_value}")
        text = lable.text()
        value = normalize_number(new_value)
        _index = text.index(":") + 1
        text = f"{text[:_index]} {value} {size}"
        lable.setText(text)

    def recount_mass(self):
        dialog = QtWidgets.QInputDialog()
        new_summ, ok = dialog.getDouble(self, "Пересчитать массу на новую",
                                          "Новая масса:", 100.00, min=1.0, decimals=2, step=5)
        if ok and new_summ:
            summ = Decimal(0)
            new_summ = Decimal(new_summ)
            list_obj = self.component_one.get_list_obj()
            if self.recepture_data.flag_2k:
                list_obj += self.component_two.get_list_obj()

            valid_list_obj = []
            for row in list_obj:
                row: ComponentRow
                row_data = row.get_data()
                if len(row_data) == 2:
                    mass = row_data[1].replace(",",".").strip()
                    if mass == "" or mass == ".":
                        continue
                    summ += Decimal(mass)
                    valid_list_obj.append(row)

            for row in valid_list_obj:
                row_data = row.get_data()
                mass = Decimal(row_data[1].replace(",", ".").strip())

                new_mass = (mass / summ) * new_summ
                row.set_amount(normalize_number(new_mass))

    def add_solvent(self):
        dialog = QtWidgets.QInputDialog()
        goal_mass, ok = dialog.getDouble(self, "Довести растворителем до массы",
                                        "Новая масса:", 100.00, min=1.0, decimals=2, step=5)
        if ok and goal_mass:
            goal_mass = Decimal(goal_mass)
            mass_suhoi = Decimal(0)
            mass_solvent = Decimal(0)
            chech_solvent_exist = False
            amount_solvents = Decimal(0)
            list_obj = self.component_one.get_list_obj()
            if self.recepture_data.flag_2k:
                list_obj += self.component_two.get_list_obj()

            solvent_list_obj = []
            for row in list_obj:
                row: ComponentRow
                row_data = row.get_data()
                if len(row_data) == 2:
                    mass = row_data[1].replace(",", ".").strip()
                    category = row.get_category()
                    if (mass == "" or mass == "." or category == "") and category != 'Solvents':
                        continue
                    if category == 'Solvents':
                        if mass == "" or mass == ".":
                            mass = "0"
                            row.set_amount(mass)
                        amount_solvents += 1
                        chech_solvent_exist = True
                        mass_solvent += Decimal(mass.replace(",", "."))
                        solvent_list_obj.append(row)
                    else:
                        mass_suhoi += Decimal(mass.replace(",", "."))

            need_solvent = goal_mass - mass_suhoi
            if need_solvent < 0 or not chech_solvent_exist:
                InfoWindow("В рецептуре не указаны растворители,\nлибо желаемая масса меньше возможной.").exec()
                return

            for row in solvent_list_obj:
                row_data = row.get_data()
                mass = Decimal(row_data[1].replace(",", ".").strip())

                if mass_solvent != 0:
                    mass = (mass * need_solvent) / mass_solvent
                else:
                    mass = need_solvent / amount_solvents

                row.set_amount(normalize_number(mass))

    def subtract_warehouse(self):
        if InfoWindow(f"Вычесть массы компонентов со склада?").exec():
            list_error_comp = []
            warehouse_error_check = False

            list_obj = self.component_one.get_list_obj()
            if self.recepture_data.flag_2k:
                list_obj += self.component_two.get_list_obj()

            actual_list_obj = []
            for row in list_obj:
                row: ComponentRow
                row_data = row.get_data()
                if len(row_data) == 2:
                    mass = row_data[1].replace(",", ".").strip()
                    name = row_data[0]
                    category = row.get_category()
                    if category != "":
                        if mass == '' or mass == ".":
                            mass = "0"
                            row.set_amount(mass)
                        mass = Decimal(mass)
                        warehouse = self.db.get_info_reactive(category, name, 'warehouse')[0][0]
                        warehouse = warehouse.replace(",", ".").strip()
                        if warehouse == '' or warehouse == ".":
                            warehouse = "0"
                        warehouse = Decimal(warehouse)
                        deffer = warehouse - mass
                        if deffer < 0:
                            warehouse_error_check = True
                            list_error_comp.append(name)
                        actual_list_obj.append((row, warehouse))

            if warehouse_error_check:
                str_comp = ",\n".join(list_error_comp)
                InfoWindow(f"Массы не вычтены. Не хватает \nследующих компонентов:\n{str_comp}").exec()
                return

            for row, warehouse in actual_list_obj:
                row: ComponentRow
                row_data = row.get_data()

                mass = Decimal(row_data[1].replace(",", ".").strip())
                name = row_data[0]
                category = row.get_category()

                deffer = warehouse - mass
                deffer = normalize_number(deffer)
                self.db.update_warehouse(category, name, deffer)
            InfoWindow("Остатки реактивов изменены.").exec()

    def open_r_settings(self):
        if self.settings_window is None:
            self.settings_window = ReceptureSettings(self)
            self.settings_window.show()

    def open_count_additives(self):
        if self.additive_window is None:
            self.collect_rows_data()
            self.additive_window = CountAdditiveWindow(self)
            self.additive_window.show()

    def open_count_hardeners(self):
        if self.hardener_window is None:
            self.collect_rows_data()
            self.hardener_window = CountHardenerWindow(self)
            self.hardener_window.show()

    def open_recount_on_maslo(self):
        if self.recount_on_maslo_window is None:
            self.collect_rows_data()
            self.recount_on_maslo_window = RecountOnMaslo(self)
            self.recount_on_maslo_window.show()

    def open_count_recepture_const(self):
        if self.count_recepture_constant is None:
            self.collect_rows_data()
            self.recepture_data.create_category_objs()
            self.recepture_data.count_mass()
            self.recepture_data.all_mass_for_suhoi_f()
            self.recepture_data.all_suhoi_in_objects()
            self.recepture_data.all_oil_in_objects()
            self.count_recepture_constant = CountReceptureConstant(self)
            self.count_recepture_constant.show()

    def open_count_recepture_comb(self):
        if self.count_recepture_combo is None:
            self.collect_rows_data()
            self.recepture_data.create_category_objs()
            self.recepture_data.count_mass()
            self.recepture_data.all_mass_for_suhoi_f()
            self.recepture_data.all_suhoi_in_objects()
            self.recepture_data.all_oil_in_objects()
            self.recepture_data.all_hiding_in_objects()
            self.count_recepture_combo = CountReceptureCombo(self)
            self.count_recepture_combo.show()

    def open_philums(self):
        if self.philum_window is None:
            self.philum_window = PhilumWindow(self)
            self.philum_window.show()

    def open_choice_color(self):
        if self.choice_color_window is None:
            self.choice_color_window = ChoiceColor(self, self.set_selected_color)
            self.choice_color_window.show()

    def set_selected_color(self, argb):
        image = QtGui.QPixmap(generate_color(argb))
        self.color2.setPixmap(image)
        self.recepture_data.recepture_color = argb

        delta_e = self.recepture_data.count_delta_e_color()
        self.delta_e_color_l.setText(f"ΔЕ: {delta_e}")


class ComponentRow(QtWidgets.QFrame):
    def __init__(self, parent, _index, name="", amount="", callback_get_list_obj=None, callback_mass=None,
                 comment="$None"):
        super(ComponentRow, self).__init__(parent=parent)
        self.callback_mass = callback_mass
        self.db = DB()
        self.category = ""
        self.category_obj = None
        self.callback_get_list_obj = callback_get_list_obj
        self.flag_comment = False
        self.component_list = []
        self.event_list = []

        self.setCursor(QtGui.QCursor(QtCore.Qt.CursorShape.SizeAllCursor))
        self.horizontalLayout_4 = QtWidgets.QHBoxLayout(self)
        self.horizontalLayout_4.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout_4.setSpacing(3)

        self.category_icon = QtWidgets.QLabel(parent=self)
        self.category_icon.setMaximumSize(QtCore.QSize(16, 16))
        self.category_icon.setMinimumSize(QtCore.QSize(16, 16))
        self.category_icon.setTextFormat(QtCore.Qt.TextFormat.PlainText)
        self.category_icon.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.category_icon.setScaledContents(True)
        self.horizontalLayout_4.addWidget(self.category_icon)
        self.assign_category(name)

        self.number_l = QtWidgets.QLabel(parent=self)
        self.number_l.setMaximumSize(QtCore.QSize(16, 16))
        self.number_l.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.number_l.setMinimumSize(QtCore.QSize(16, 16))
        self.horizontalLayout_4.addWidget(self.number_l)

        list_all_names = self.db.search("%%")
        self.name_comp = SearchCombobox(self, list_all_names)
        self.name_comp.setText(name)
        self.name_comp.setMinimumSize(QtCore.QSize(250, 0))
        self.name_comp.textChanged.connect(self.name_changed)


        self.horizontalLayout_4.addWidget(self.name_comp)
        self.amount = CustomEntry(self, padding=False)
        self.amount.setText(amount)
        self.amount.setMaximumSize(QtCore.QSize(50, 16777215))
        self.amount.setMinimumSize(QtCore.QSize(50, 16777215))
        self.amount.textChanged.connect(lambda event: callback_mass())
        self.amount.setValidator(get_numeric_validator())
        self.horizontalLayout_4.addWidget(self.amount)

        self.comment_spacer = QtWidgets.QSpacerItem(40, 10, QtWidgets.QSizePolicy.Policy.Fixed,
                                            QtWidgets.QSizePolicy.Policy.Minimum)
        self.comment_spacer.changeSize(0, 0)
        self.horizontalLayout_4.addItem(self.comment_spacer)
        self.comment = QtWidgets.QPlainTextEdit(parent=self)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Policy.Minimum, QtWidgets.QSizePolicy.Policy.Minimum)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(self.comment.sizePolicy().hasHeightForWidth())
        self.comment.setSizePolicy(sizePolicy)
        # self.comment.setContentsMargins(32,0,0,0)
        self.comment.setMinimumSize(QtCore.QSize(303, 40))
        self.horizontalLayout_4.addWidget(self.comment)
        self.comment.hide()

        self.swap = MenuButton(self, "swap", (20, 20))
        self.swap.setMaximumSize(QtCore.QSize(20, 20))

        menu = CustomMenu(self)
        menu.addAction('Сделать комментарием', self.change_state)
        menu.addAction('Удалить', self.delete)

        self.swap.setMenu(menu)
        self.horizontalLayout_4.addWidget(self.swap)

        if comment != "$None":
            self.change_state(initial=True)
            self.comment.setPlainText(comment)

    def change_state(self, initial=False):
        if self.flag_comment:
            self.flag_comment = False
            self.category_icon.show()
            self.number_l.show()
            self.name_comp.show()
            self.amount.show()
            self.assign_category(self.name_comp.text())

            self.comment.hide()
            self.comment_spacer.changeSize(0, 0)
            menu = QtWidgets.QMenu(self)
            menu.addAction('Сделать комментарием', self.change_state)
            menu.addAction('Удалить', self.delete)

            self.swap.setMenu(menu)
            self.reset_row_number()

        else:
            self.flag_comment = True
            self.category = ""
            self.category_icon.hide()
            self.number_l.hide()
            self.name_comp.hide()
            self.amount.hide()

            self.comment.show()
            self.comment_spacer.changeSize(38, 10)
            menu = QtWidgets.QMenu(self)
            menu.addAction('Сделать компонентом', self.change_state)
            menu.addAction('Удалить', self.delete)

            self.swap.setMenu(menu)
            self.reset_row_number()
        if not initial:
            self.callback_mass()

        menu.setStyleSheet(
            """
            QMenu
            {
                font: 10pt;
                background-color: #eee;
            }
            QMenu::item:selected
            {
                background-color: #209fa6
            }
            """
            )

    def assign_category(self, name):
        text = ""
        tooltip = ""
        if len(self.db.check_group_reactives("Solvents", name)) == 1:
            self.category_obj = Solvents(name, "0")
            self.category = "Solvents"
            icon_path = "images/solvent.png"
            tooltip = self.category_obj.to_string()
        elif len(self.db.check_group_reactives("Pigments", name)) == 1:
            self.category_obj = Pigments(name, "0")
            self.category = "Pigments"
            icon_path = "images/pigment.png"
            tooltip = self.category_obj.to_string()
        elif len(self.db.check_group_reactives("Fillers", name)) == 1:
            self.category_obj = Fillers(name, "0")
            self.category = "Fillers"
            icon_path = "images/filler.png"
            tooltip = self.category_obj.to_string()
        elif len(self.db.check_group_reactives("Films", name)) == 1:
            self.category_obj = Films(name, "0")
            self.category = "Films"
            icon_path = "images/film.png"
            tooltip = self.category_obj.to_string()
        elif len(self.db.check_group_reactives("Additives", name)) == 1:
            self.category_obj = Additives(name, "0")
            self.category = "Additives"
            icon_path = "images/additive.png"
            tooltip = self.category_obj.to_string()
        elif len(self.db.check_group_reactives("PigmPast", name)) == 1:
            self.category_obj = Pigmpasts(name, "0")
            self.category = "PigmPast"
            icon_path = "images/pigm_past.png"
            tooltip = self.category_obj.to_string()
        elif len(self.db.check_group_reactives("Hardener", name)) == 1:
            self.category_obj = Hardeners(name, "0")
            self.category = "Hardener"
            icon_path = "images/hardener.png"
            tooltip = self.category_obj.to_string()
        else:
            self.category = ""
            icon_path = ""
            text = "?"
            tooltip = "Поле пустое или компонент не найден.\n" \
                      "Данная строка не будет учитываться\n" \
                      "в расчете."

        self.category_icon.setText(text) if text != "" else self.category_icon.setPixmap(QtGui.QPixmap(icon_path))
        self.category_icon.setToolTip(tooltip)

    def set_number(self, number: int):
        self.number_l.setText(str(number))

    def delete(self, event=None):
        # self.hide()
        parent: Ui_Component = self.parent()
        print(self)
        list_obj = parent.get_list_obj(raw=True)
        _index = list_obj.index(self)
        print(list_obj)
        for_delete = parent.gridLayout.takeAt(_index)

        print(for_delete.layout().deleteLater())
        print(parent.get_list_obj(raw=True))

        self.hide()
        self.reset_row_number()
        self.name_comp.setFocus()

    def name_changed(self, event):
        self.assign_category(event)

    def reset_row_number(self):
        self.list_obj: List[ComponentRow]
        number = 1
        for obj in self.callback_get_list_obj():
            if obj is not None:
                if not obj.flag_comment:
                    obj.set_number(number)
                    number += 1

    def get_data(self):
        #если 1 элемент, то комментарий, если 2, то компонент
        if self.flag_comment:
            return self.comment.toPlainText()
        else:
            return self.name_comp.text(),  self.amount.text()

    def set_amount(self, amount: str):
        self.amount.setText(amount)

    def get_category(self):
        return self.category


class Ui_Component(QtWidgets.QWidget):

    list_obj = []

    @staticmethod
    def set_drop_false(obj = None):
        for i in Ui_Component.list_obj:
            if i is not None and i is not obj:
                i.setAcceptDrops(False)

    def __init__(self, parent, recepture_data):
        super(Ui_Component, self).__init__(parent=parent)
        self.target = None
        Ui_Component.list_obj.append(self)
        self.setAcceptDrops(False)
        self.list_comp_row_obj = []
        # self.recepture_data = recepture_data
        self.row = 0
        self.setGeometry(QtCore.QRect(0, 0, 403, 485))
        self.setObjectName("recepture")
        # self.component_one_l = QtWidgets.QVBoxLayout(self)
        # self.component_one_l.setContentsMargins(0, 0, 0, 0)
        # self.component_one_l.setSpacing(2)

        # self.layout = QtWidgets.QHBoxLayout(self)
        # self.scrollArea = QtWidgets.QScrollArea(self)
        # self.scrollArea.setWidgetResizable(True)
        # self.scrollAreaWidgetContents = QtWidgets.QWidget()
        self.gridLayout = QtWidgets.QGridLayout(self)
        self.gridLayout.setSpacing(2)
        self.gridLayout.setContentsMargins(0,0,0,0)

        # self.scrollArea.setWidget(self.scrollAreaWidgetContents)
        # self.layout.addWidget(self.scrollArea)
        # for name, value in self.recepture_data.component_list:
        #     self.add_row("one", name=name, value=value)

        # for i in range(3):
        #     for j in range(3):
        #         self.Frame = QtWidgets.QFrame(self)
        #         self.Frame.setStyleSheet("background-color: white;")
        #
        #         self.Frame.setLineWidth(2)
        #         self.layout = QtWidgets.QHBoxLayout(self.Frame)
        #         l = QtWidgets.QLabel(parent=self.Frame)
        #         l.setText(str(i) + str(j))
        #         Box = QtWidgets.QVBoxLayout()
        #
        #         Box.addWidget(self.Frame)
        #
        #         self.gridLayout.addLayout(Box, i, j)
        #         self.gridLayout.setColumnStretch(i % 3, 1)
        #         self.gridLayout.setRowStretch(j, 1)

    def add_row(self, name="", value="", callback_mass=None, comment=None):
        parent = self
        list_obj = self.list_comp_row_obj
        loyout = self.gridLayout

        _index = len(list_obj)
        row = ComponentRow(parent, _index, name=name, amount=value,
                           callback_get_list_obj=self.get_list_obj, callback_mass=callback_mass, comment=comment)

        Box = QtWidgets.QVBoxLayout()
        Box.addWidget(row)

        loyout.addLayout(Box, self.row, 0)
        # loyout.setColumnStretch(i % 3, 1)
        # loyout.setRowStretch(j, 1)

        list_obj.append(row)
        self.row += 1
        self.reset_row_number()

    def get_list_obj(self, raw=False):
        list_obj = []
        raw_list_obj = []
        for i in range(self.gridLayout.count()):
            component_row_obj: QtWidgets.QWidgetItem = self.gridLayout.itemAt(i).itemAt(0)
            if component_row_obj.widget():
                list_obj.append((component_row_obj.widget(), self.gridLayout.getItemPosition(i)[0]))
                raw_list_obj.append(component_row_obj.widget())
        if raw:
            return raw_list_obj

        list_obj.sort(key=lambda x: x[1])
        list_obj = list(map(lambda x:x[0], list_obj))
        return list_obj

    def reset_row_number(self):
        list_obj: List[ComponentRow]
        list_obj = self.get_list_obj()
        number = 1
        for obj in list_obj:
            if obj is not None and not obj.isHidden():
                if not obj.flag_comment:
                    obj.set_number(number)
                    number += 1


    def get_index(self, pos):
        for i in range(self.gridLayout.count()):
            if self.gridLayout.itemAt(i).geometry().contains(pos) and i != self.target:
                return i

    def mousePressEvent(self, event: QtGui.QMouseEvent):
        if event.button() == Qt.MouseButton.LeftButton:
            self.setAcceptDrops(True)
            self.target = self.get_index(event.position().toPoint())
            Ui_Component.set_drop_false(obj=self)
        else:
            self.target = None
            Ui_Component.set_drop_false()

    def mouseMoveEvent(self, event):
        if event.buttons() & Qt.MouseButton.LeftButton and self.target is not None:
            drag = QtGui.QDrag(self.gridLayout.itemAt(self.target))
            pix = self.gridLayout.itemAt(self.target).itemAt(0).widget().grab()
            mimedata = QtCore.QMimeData()
            mimedata.setImageData(pix)
            drag.setMimeData(mimedata)
            drag.setPixmap(pix)
            # drag.setHotSpot(event.pos())
            drag.exec()

    def mouseReleaseEvent(self, event):
        self.target = None
        Ui_Component.set_drop_false()

    def dragEnterEvent(self, event):
        if event.mimeData().hasImage():
            event.accept()
        else:
            event.ignore()

    def dropEvent(self, event: QtGui.QMouseEvent):
        test:QtWidgets.QVBoxLayout = event.source()
        print(test.parentWidget())
        if not event.source().geometry().contains(event.position().toPoint()):
            source = self.get_index(event.position().toPoint())
            if source is None:
                return
            list_obj = self.get_list_obj()
            raw_list_obj = self.get_list_obj(raw=True)
            s = list_obj.index(self.gridLayout.itemAt(source).itemAt(0).widget())
            f = list_obj.index(self.gridLayout.itemAt(self.target).itemAt(0).widget())

            if s > f:
                for i in range(f + 1, s + 1):
                    raw_index = self.get_list_obj(raw=True).index(list_obj[i])
                    p1 = list(self.gridLayout.getItemPosition(raw_index))
                    p1[0] = i - 1
                    self.gridLayout.addItem(self.gridLayout.takeAt(raw_index), *p1)
                raw_index = self.get_list_obj(raw=True).index(list_obj[f])
                self.gridLayout.addItem(self.gridLayout.takeAt(raw_index), s, 0, 1, 1)
            elif s < f:
                for i in range(s, f):
                    raw_index = self.get_list_obj(raw=True).index(list_obj[i])
                    p1 = list(self.gridLayout.getItemPosition(raw_index))
                    p1[0] = i + 1
                    self.gridLayout.addItem(self.gridLayout.takeAt(raw_index), *p1)
                raw_index = self.get_list_obj(raw=True).index(list_obj[f])
                self.gridLayout.addItem(self.gridLayout.takeAt(raw_index), s, 0, 1, 1)


            # i, j = max(self.target, source), min(self.target, source)
            # p1, p2 = self.gridLayout.getItemPosition(i), self.gridLayout.getItemPosition(j)
            # print(p1)
            # self.gridLayout.addItem(self.gridLayout.takeAt(i), *p2)
            # self.gridLayout.addItem(self.gridLayout.takeAt(j), *p1)
            self.reset_row_number()
        Ui_Component.set_drop_false()


class ExperimentRow:
    row = 2

    def __init__(self, parent, lo: QtWidgets.QGridLayout, name: str, needed: str, value: str, state: int):
        self.name = name
        self.needed = needed
        name_l = QtWidgets.QLabel(parent=parent)
        name_l.setMinimumSize(QtCore.QSize(200, 0))
        name_l.setText(name)
        lo.addWidget(name_l, ExperimentRow.row, 0, 1, 1)

        needed_l = QtWidgets.QLabel(parent=parent)
        needed_l.setMaximumSize(QtCore.QSize(100, 16777215))
        needed_l.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        needed_l.setText(needed)
        lo.addWidget(needed_l, ExperimentRow.row, 1, 1, 1)

        self.value = CustomEntry(parent=parent)
        self.value.setMaximumSize(QtCore.QSize(100, 16777215))
        self.value.setText(value)
        lo.addWidget(self.value, ExperimentRow.row, 2, 1, 1)

        number_group = QtWidgets.QButtonGroup(parent)
        self.gray_rb = CustomRadioBtn("grey")
        self.gray_rb.setText("")
        self.gray_rb.setObjectName("gray_rb")
        number_group.addButton(self.gray_rb)
        lo.addWidget(self.gray_rb, ExperimentRow.row, 4, 1, 1)

        self.green_rb = CustomRadioBtn("green")
        self.green_rb.setText("")
        self.green_rb.setObjectName("green_rb")
        number_group.addButton(self.green_rb)
        lo.addWidget(self.green_rb, ExperimentRow.row, 3, 1, 1)

        self.red_rb = CustomRadioBtn("red")
        self.red_rb.setText("")
        self.red_rb.setObjectName("red_rb")
        number_group.addButton(self.red_rb)
        lo.addWidget(self.red_rb, ExperimentRow.row, 5, 1, 1)

        if int(state) == 0:
            self.gray_rb.setChecked(True)
        elif int(state) == 1:
            self.green_rb.setChecked(True)
        elif int(state) == -1:
            self.red_rb.setChecked(True)

        ExperimentRow.row += 1

    def get(self):
        state = 0
        if self.green_rb.isChecked():
            state = 1
        if self.red_rb.isChecked():
            state = -1
        # print(state)
        return self.name, self.needed, self.value.text(), state


class ReceptureDataModel:
    def __init__(self, project, iteration, name):
        self.project = project
        self.iteration = iteration
        self.name = name
        self.not_encoded_projects = ['Тарировочные_кривые', 'Тарировочные кривые', 'Примеры']

        self.project_params = []
        self.project_params_value = []
        self.project_color = "#00ffffff"
        self.recepture_color = "#00ffffff"
        self.password = ""

        self.component_list = [("", "") for _ in range(7)]
        self.list_comments = ["$None" for _ in range(7)]
        self.category_list = ["" for _ in range(7)]

        self.component_list_2 = [("", "") for _ in range(3)]
        self.list_comments_2 = ["$None" for _ in range(3)]
        self.category_list_2 = ["" for _ in range(3)]

        self.experiment_list = []
        self.list_experiment_status = None
        self.notes = ""

        self.flag_2k = False
        self.price_K = 1.0
        self.accurate_density = 0.0

        self.mass = Decimal(0)
        self.price = Decimal(0)
        self.suhoi = Decimal(0)
        self.degree_pigm = Decimal(0)
        self.oil = Decimal(0)
        self.const_pigm = Decimal(0)
        self.hiding_pigm = Decimal(0)
        self.hiding_wet = Decimal(0)
        self.hiding_dry = Decimal(0)
        self.philum = Decimal(0)
        self.okp = Decimal(0)
        self.kokp = Decimal(0)
        self.okp_kokp = Decimal(0)
        self.hiding_dry = Decimal(0)
        self.density = Decimal(0)
        self.volume_suhoi = Decimal(0)

    def get_density(self):
        if self.accurate_density > 0.0:
            return Decimal(self.accurate_density)
        else:
            return self.density

    def map_encrypt(self, str):
        global password
        if self.project not in self.not_encoded_projects and str != "$None":
            result = Secrets().symmetric_encrypt(str.encode(), password)
        else:
            result = str
        return result

    def map_decrypt(self, byte):
        if self.project not in self.not_encoded_projects and byte != "$None":
            result = Secrets().symmetric_decrypt(byte, password).decode()
        else:
            result = byte
        return result

#TODO: Сохранять категории и загружать
    def load_data(self):
        with SqliteDict('saves/' + self.project + '/params') as mydict:
            enc_data_params = mydict['params']
            enc_data_params_value = mydict['params_value']
            if self.project not in self.not_encoded_projects:
                for params, params_value in zip(enc_data_params, enc_data_params_value):
                    self.project_params.append(Secrets().symmetric_decrypt(params, self.password).decode())
                    self.project_params_value.append(Secrets().symmetric_decrypt(params_value, self.password).decode())
            else:
                self.project_params = enc_data_params
                self.project_params_value = enc_data_params_value
            self.project_color = mydict.get("project_color", "#00ffffff")


        with SqliteDict('saves/' + self.project + '/' + self.iteration) as mydict:
            data_iteraton = dict(mydict)
            self.data = data_iteraton.get(self.name, None)

            # [0] - реактивы, [1]- масса реактивов, [2] - эксперимент параметры,
            # [3] - полученны значения, [4] - заметки, [5] - ТЗ, [6] - расчетные характеристики,
            # [7] - реактивы 2к, [8] - масса реактивов 2к, [9] - dict params
        if self.data is None:
            empty_exp_value = ("" for _ in self.project_params)
            empty_exp_status = (0 for _ in self.project_params)
            self.experiment_list = list(zip(self.project_params, self.project_params_value, empty_exp_value, empty_exp_status))

            return

        if len(self.data) > 9:
            configs = self.data.pop(9)
            self.price_K = configs.get('price_K', 1.0)
            self.accurate_density = configs.get('accurate_density', 0.0)
            self.list_experiment_status = configs.get('list_experiment_status', [0 for _ in self.data[2]])
            self.recepture_color = configs.get("recepture_color", "#00ffffff")
            list_comments = configs.get("list_comments", ["$None" for _ in self.data[0]])
            list_comments_2 = configs.get("list_comments_2", ["$None" for _ in self.data[0]])
            self.list_comments = list(map(self.map_decrypt, list_comments))
            self.list_comments_2 = list(map(self.map_decrypt, list_comments_2))
            self.okp_kokp = configs.get("okp_kokp", Decimal(0))
            self.volume_suhoi = configs.get("volume_suhoi", Decimal(0))
        else:
            self.price_K = 1.0
            self.accurate_density = 0.0
            self.list_experiment_status = [0 for _ in self.data[2]]
            self.recepture_color = "#00ffffff"
            self.list_comments = ["$None" for _ in self.data[0]]
            self.list_comments_2 = ["$None" for _ in self.data[0]]
            self.okp_kokp = Decimal(0)
            self.volume_suhoi = Decimal(0)

        for i, param in enumerate(self.data):
            self.data[i] = list(map(self.map_decrypt, param))

        self.component_list = list(zip(self.data[0], self.data[1]))
        self.component_list_2 = list(zip(self.data[7], self.data[8]))
        self.experiment_list = list(zip(self.data[2], self.data[5], self.data[3], self.list_experiment_status))
        notes = self.data[4]
        if isinstance(notes, list):
            notes = notes[0]
        self.notes = notes

        # self.properies=['Цена','м.д.н.в','СП','ОКП','Масло','Кп','Ср укрыв','Укрыв сух пленки','Филум', 'КОКП', 'Укр мокрой пл', 'плотность']
        properties = self.data[6]
        self.price = properties[0]
        self.suhoi = properties[1]
        self.degree_pigm = properties[2]
        self.okp = properties[3]
        self.oil = properties[4]
        self.const_pigm = properties[5]
        self.hiding_pigm = properties[6]
        self.hiding_wet = properties[7]
        self.philum = properties[8]
        self.kokp = properties[9]
        self.hiding_dry = properties[10]
        self.density = properties[11]

        for comp in self.component_list_2:
            if comp[0] != '' and comp[1] != '':
                self.flag_2k = True

    def save(self, event=None, save_as=None):
        if save_as is not None:
            self.iteration = save_as[0]
            self.name = save_as[1]

        self.all_count()
        # [0] - реактивы, [1]- масса реактивов, [2] - эксперимент параметры,
        # [3] - полученны значения, [4] - заметки, [5] - ТЗ, [6] - расчетные характеристики,
        # [7] - реактивы 2к, [8] - масса реактивов 2к, [9] - dict params
        reactives = []
        reactives_mass = []
        experiment_params = []
        experiment_value = []
        note = self.notes
        needed_experiment_value = []
        properies = []
        reactives_2 = []
        reactives_mass_2 = []
        dict_params = {}
        list_comments = []
        list_comments_2 = []
        list_experiment_status = []
        # print(self.component_list)
        for row in self.component_list:
            if isinstance(row, str):
                reactives.append("")
                reactives_mass.append("")
                list_comments.append(row)
            else:
                reactives.append(row[0])
                reactives_mass.append(row[1])
                list_comments.append("$None")

        for row in self.component_list_2:
            if isinstance(row, str):
                reactives_2.append("")
                reactives_mass_2.append("")
                list_comments_2.append(row)
            else:
                reactives_2.append(row[0])
                reactives_mass_2.append(row[1])
                list_comments_2.append("$None")

        for name, need, value, status in self.experiment_list:
            experiment_params.append(name)
            experiment_value.append(value)
            needed_experiment_value.append(need)
            list_experiment_status.append(status)

        # self.properies=['Цена','м.д.н.в','СП','ОКП','Масло','Кп','Ср укрыв','Укрыв сух пленки','Филум', 'КОКП', 'Укр мокрой пл', 'плотность']
        # old style
        properies = [
            self.price,
            self.suhoi,
            self.degree_pigm,
            self.okp,
            self.oil,
            self.const_pigm,
            self.hiding_pigm,
            self.hiding_dry,
            self.philum,
            self.kokp,
            self.hiding_wet,
            self.get_density(),
        ]

        dict_params['price_K'] = self.price_K
        dict_params['accurate_density'] = self.accurate_density
        dict_params["list_experiment_status"] = list_experiment_status
        dict_params["recepture_color"] = self.recepture_color
        dict_params["list_comments"] = list(map(self.map_encrypt, list_comments))
        dict_params["list_comments_2"] = list(map(self.map_encrypt, list_comments_2))
        dict_params["okp_kokp"] = self.okp_kokp
        dict_params["volume_suhoi"] = self.volume_suhoi

        with SqliteDict('saves/' + self.project + '/' + self.iteration) as mydict:

            # [0] - реактивы, [1]- масса реактивов, [2] - эксперимент параметры,
            # [3] - полученны значения, [4] - заметки, [5] - ТЗ, [6] - расчетные характеристики,
            # [7] - реактивы 2к, [8] - масса реактивов 2к, [9] - dict params
            mydict[self.name] = [list(map(self.map_encrypt, reactives)),
                                 list(map(self.map_encrypt, reactives_mass)),
                                 list(map(self.map_encrypt, experiment_params)),
                                 list(map(self.map_encrypt, experiment_value)),
                                 list(map(self.map_encrypt, [note, ])),
                                 list(map(self.map_encrypt, needed_experiment_value)),
                                 list(map(self.map_encrypt, properies)),
                                 list(map(self.map_encrypt, reactives_2)),
                                 list(map(self.map_encrypt, reactives_mass_2)),
                                 dict_params,
                                 ]
            mydict.commit()

    def delete(self):
        with SqliteDict('saves/' + self.project + '/' + self.iteration) as mydict:
            mydict.pop(self.name)
            mydict.commit()

    def get_count_dict(self) -> dict:
        recepture = self
        map_dict_param = {
            "Цена": recepture.price,
            "Маслоемкость 1-го рода": recepture.oil,
            "Масс.д.н.в": recepture.suhoi,
            "Объем.д.н.в": recepture.volume_suhoi,
            "ОКП": recepture.okp,
            "КОКП": recepture.kokp,
            "ОКП/КОКП": recepture.okp_kokp,
            "Укрывистость пигментов": recepture.hiding_pigm,
            "Укрывистость мокрой пленки": recepture.hiding_wet,
            "Укрывистость сухой пленки": recepture.hiding_dry,
            "Филум пигментов": recepture.philum,
            "Плотность": recepture.get_density(),
            "Степень пигментирования": recepture.degree_pigm,
            "Константа наполнения": recepture.const_pigm,
        }
        return map_dict_param

    def count_mass(self, all=False) -> Decimal:
        components, _ = self.get_actual_data_for_count(all=all)
        components = list(map(lambda x: Decimal(x[1].replace(",", ".")), components))
        if len(components) !=0:
            summ = reduce(lambda x, y: x + y, components)
        else:
            summ = Decimal(0)
        self.mass = summ
        return self.mass

    def get_actual_data_for_count(self, all=False) -> Tuple[tuple, str]:
        # retrun [((name, value), category), ...]
        # retrun [((comment,), category), ...]

        if self.flag_2k:
            components = self.component_list + self.component_list_2
            categories = self.category_list + self.category_list_2
        else:
            components = self.component_list
            categories = self.category_list

        data = list(zip(components, categories))
        data = list(filter(lambda foo: len(foo[0]) > 1, data))
        if all:
            data = list(filter(lambda foo: len(foo[0][1]) > 0, data))
            # print(data)

        else:
            data = list(filter(lambda foo: foo[1] != "", data))

        data = list(filter(lambda foo: type(foo[0]) == tuple, data))
        if len(list(zip(*data))) != 0:
            components, categories = list(zip(*data))
        else:
            components, categories = [], []
            print("Нет компонентов в рецептуре")
        return components, categories

    def create_category_objs(self):
        components, categories = self.get_actual_data_for_count()

        self.list_of_solvent_objects = []
        self.list_of_pigment_objects = []
        self.list_of_filler_objects = []
        self.list_of_film_objects = []
        self.list_of_additive_objects = []
        self.list_of_pigmpast_objects = []
        self.list_of_hardener_objects = []

        for (name, mass), cat in zip(components, categories):
            if cat == "Solvents":
                self.list_of_solvent_objects.append(Solvents(name, mass))
            elif cat == "Pigments":
                self.list_of_pigment_objects.append(Pigments(name, mass))
            elif cat == "Fillers":
                self.list_of_filler_objects.append(Fillers(name, mass))
            elif cat == "Films":
                self.list_of_film_objects.append(Films(name, mass))
            elif cat == "Additives":
                self.list_of_additive_objects.append(Additives(name, mass))
            elif cat == "PigmPast":
                self.list_of_pigmpast_objects.append(Pigmpasts(name, mass))
            elif cat == "Hardener":
                self.list_of_hardener_objects.append(Hardeners(name, mass))
            else:
                pass

    def get_components_obj(self):
        self.create_category_objs()

        return self.list_of_solvent_objects + self.list_of_pigment_objects + self.list_of_filler_objects + \
        self.list_of_film_objects + self.list_of_additive_objects + self.list_of_pigmpast_objects +\
        self.list_of_hardener_objects

    def all_count(self):
        info_text = ""
        self.count_mass()
        self.create_category_objs()

        list_count_funcs = [(self.count_price, "Ошибка в расчете цены. Переведите валюту в Руб."),
                            (self.all_mass_for_suhoi_f, "Ошибка в расчете массы"),
                            (self.all_suhoi_in_objects, "Ошибка при расчете м.д.н.в."),
                            (lambda: self.all_suhoi_in_objects(for_volume=True), "Ошибка при расчете м.д.н.в."),
                            (self.all_density_in_objects, "Ошибка в расчете плотности"),
                            (self.all_degree_pigm_in_objects, "Ошибка в расчете степени пигм."),
                            (self.all_okp_in_objects, "Ошибка в расчете ОКП"),
                            (self.all_oil_in_objects, "Ошибка в расчете маслоемкости"),
                            (self.all_const_pigm_in_objects, "Ошибка в расчете кН"),
                            (self.all_hiding_in_objects, "Ошибка в расчете укрывистости"),
                            (self.filum_in_objects, "Ошибка в расчете филума"),
                            (self.all_kokp_maslo_in_objects, "Ошибка в расчете КОКП"),
                            (self.all_volume_suhoi_in_objects, 'Ошибка в расчете об. доли нелетучих в-в'),
                            ]


        for count_func, error_text in list_count_funcs:
            try:
                count_func()
            except Exception as e:
                # logging.error(e, exc_info=True)
                info_text = info_text + error_text + "\n"
                print(traceback.format_exc())
                # raise e

        if info_text != "":
            info_text = 'ОШИБКА ПРИ РАСЧЕТЕ! \nПроверьте расчетные значения используемых ' \
                        '\nкомпонентов в Моя лаборатория.\n' + info_text
            InfoWindow(info_text).exec()

    def count_price(self):
        all_mass = self.mass
        all_price = Decimal(0)

        for i in self.list_of_solvent_objects:
            all_price += (i.mass * i.price) / all_mass
        for i in self.list_of_pigment_objects:
            all_price += (i.mass * i.price) / all_mass
        for i in self.list_of_filler_objects:
            all_price += (i.mass * i.price) / all_mass
        for i in self.list_of_film_objects:
            all_price += (i.mass * i.price) / all_mass
        for i in self.list_of_additive_objects:
            all_price += (i.mass * i.price) / all_mass
        for i in self.list_of_pigmpast_objects:
            all_price += (i.mass * i.price) / all_mass
        for i in self.list_of_hardener_objects:
            all_price += (i.mass * i.price) / all_mass
        all_price *= Decimal(self.price_K)

        self.price = all_price


    def all_mass_for_suhoi_f(self):
        all_mass = Decimal(0)
        suhoi_check = int(get_suhoi_type())

        for i in self.list_of_solvent_objects:
            i = i.mass
            all_mass += i
        for i in self.list_of_pigment_objects:
            i = i.mass
            all_mass += i
        for i in self.list_of_filler_objects:
            i = i.mass
            all_mass += i
        for i in self.list_of_film_objects:
            i = i.mass
            all_mass += i
        for i in self.list_of_pigmpast_objects:
            i = i.mass
            all_mass += i

        if suhoi_check == 1:
            for i in self.list_of_additive_objects:
                all_mass += i.mass
        else:
            for i in self.list_of_additive_objects:
                if i.type.lower() == 'пластификатор':
                    all_mass += i.mass
        if suhoi_check == 1 or suhoi_check == 2:
            for i in self.list_of_hardener_objects:
                all_mass += i.mass

        self.all_mass_for_suhoi = all_mass

    def all_suhoi_in_objects(self, for_volume=False):
        if for_volume:
            all_mass = self.mass
        else:
            all_mass = self.all_mass_for_suhoi

        all_suhoi = Decimal(0)

        if for_volume:
            suhoi_check = 1  # учитывать всё
        else:
            suhoi_check = int(get_suhoi_type())

        for i in self.list_of_solvent_objects:
            all_suhoi += (i.mass * i.suhoi) / all_mass

        for i in self.list_of_pigment_objects:
            all_suhoi += i.mass / all_mass

        for i in self.list_of_filler_objects:
            all_suhoi += i.mass / all_mass

        for i in self.list_of_film_objects:
            all_suhoi += (i.mass * i.suhoi) / all_mass

        for i in self.list_of_additive_objects:
            if suhoi_check == 1:
                all_suhoi += (i.mass * i.suhoi) / all_mass
            if suhoi_check == 2 or suhoi_check == 3:
                if i.type.lower() == 'пластификатор':
                    all_suhoi += (i.mass * i.suhoi) / all_mass

        for i in self.list_of_pigmpast_objects:
            all_suhoi += (i.mass * i.suhoi) / all_mass

        if suhoi_check == 1 or suhoi_check == 2:
            for i in self.list_of_hardener_objects:
                all_suhoi += (i.mass * i.suhoi) / all_mass

        if not for_volume:
            all_suhoi = all_suhoi * 100
            # all_suhoi = self.normalize_number(all_suhoi)
            self.suhoi = all_suhoi
        else:
            self.all_suhoi_for_volume = all_suhoi

    def all_density_in_objects(self):
        all_mass = self.mass
        density = Decimal(0)
        for i in self.list_of_solvent_objects:
            i = i.mass * i.density / all_mass
            density += i
        for i in self.list_of_pigment_objects:
            i = i.mass * i.density / all_mass
            density += i
        for i in self.list_of_filler_objects:
            i = i.mass * i.density / all_mass
            density += i
        for i in self.list_of_film_objects:
            i = i.mass * i.density / all_mass
            density += i
        for i in self.list_of_pigmpast_objects:
            i = i.mass * i.density / all_mass
            density += i
        for i in self.list_of_additive_objects:
            i = i.mass * i.density / all_mass
            density += i
        for i in self.list_of_hardener_objects:
            i = i.mass * i.density / all_mass
            density += i

        self.density = density

    def all_volume_suhoi_in_objects(self):
        suhoi = self.all_suhoi_for_volume
        if self.accurate_density > 0.0:
            lkm_volume = Decimal(100) / Decimal(self.accurate_density)
        else:
            if self.density > 0:
                lkm_volume = Decimal(100) / Decimal(self.density)
            else:
                lkm_volume = 0
        # print('лкм объем:'+str(lkm_volume))

        all_solvent_mass = Decimal(0)
        list_same = self.list_of_film_objects + self.list_of_hardener_objects + \
                    self.list_of_pigmpast_objects + self.list_of_solvent_objects
        for i in list_same:
            all_solvent_mass += i.mass * (Decimal(1.0) - i.suhoi)

        for i in self.list_of_additive_objects:
            if i.density_solvent > 0:
                all_solvent_mass += i.mass * (Decimal(1.0) - i.suhoi)

        # print('масса растворителя:' + str(all_solvent_mass))

        density_solvent = Decimal(0)
        if all_solvent_mass != Decimal(0):
            for i in list_same:
                i = i.mass * (Decimal(1.0) - i.suhoi) * i.density_solvent / all_solvent_mass
                density_solvent += i
            for i in self.list_of_additive_objects:
                if i.density_solvent > 0:
                    i = i.mass * (Decimal(1.0) - i.suhoi) * i.density_solvent / all_solvent_mass
                    density_solvent += i
        else:
            self.volume_suhoi = Decimal(0)
            return

        # print('плотность растворителя:' + str(density_solvent))

        if density_solvent != Decimal(0) and lkm_volume != Decimal(0):
            solvent_volume = (Decimal(1.0) - suhoi) * 100 / density_solvent
            # print('объем растворителя:' + str(solvent_volume))
            volume_suhoi = (lkm_volume - solvent_volume) * 100 / lkm_volume

        else:
            volume_suhoi = Decimal(0)

        self.volume_suhoi = volume_suhoi

    def all_degree_pigm_in_objects(self):
        pigm = Decimal(0)
        film = Decimal(0)
        for i in self.list_of_pigment_objects:
            pigm += i.mass
        for i in self.list_of_filler_objects:
            pigm += i.mass
        for i in self.list_of_pigmpast_objects:
            pigm += (i.mass * i.suhoi_pigm)

        for i in self.list_of_film_objects:
            film += (i.mass * i.suhoi)
        for i in self.list_of_hardener_objects:
            film += (i.mass_for_params * i.suhoi)
        for i in self.list_of_pigmpast_objects:
            film += (i.mass * i.suhoi_film)
        for i in self.list_of_additive_objects:
            if i.type.lower() == 'пластификатор':
                film += (i.mass * i.suhoi)

        try:
            degree_pigm = Decimal(pigm / film).quantize(Decimal("1.00"), "ROUND_HALF_EVEN")
        except Exception as e:
            logging.error(e, exc_info=True)
            degree_pigm = Decimal(0)

        self.degree_pigm = degree_pigm

    def all_okp_in_objects(self):

        filler_and_pigm_volume = Decimal(0)
        film_volume = Decimal(0)

        try:
            for i in self.list_of_pigment_objects:
                filler_and_pigm_volume += (
                        (i.mass) / (i.density))
            for i in self.list_of_filler_objects:
                filler_and_pigm_volume += (
                        (i.mass) / (i.density))
            for i in self.list_of_pigmpast_objects:
                filler_and_pigm_volume += (
                        (i.mass * i.suhoi_pigm) / (
                    i.density_pigm))

            for i in self.list_of_film_objects:
                film_volume += ((i.mass * i.suhoi) / (
                    i.density_dry))
            for i in self.list_of_hardener_objects:
                film_volume += ((i.mass_for_params * i.suhoi) / (
                    i.density_dry))
            for i in self.list_of_pigmpast_objects:
                film_volume += (
                        i.mass * i.suhoi_film /
                        i.density_dry)

            for i in self.list_of_additive_objects:
                if i.type.lower() == 'пластификатор':
                    film_volume += ((i.mass * i.suhoi) / (
                        i.density))

            okp = (filler_and_pigm_volume * 100) / (filler_and_pigm_volume + film_volume)

        except Exception as e:
            logging.error(e, exc_info=True)
            okp = Decimal(0)

        self.okp = okp

    def all_oil_in_objects(self):
        all_mass = Decimal(0)
        all_maslo = Decimal(0)
        for i in self.list_of_pigment_objects:
            i = i.mass
            all_mass += i
        for i in self.list_of_filler_objects:
            i = i.mass
            all_mass += i
        for i in self.list_of_pigmpast_objects:
            i = i.mass * i.suhoi_pigm
            all_mass += i

        for i in self.list_of_pigment_objects:
            all_maslo += (i.mass * i.maslo) / all_mass

        for i in self.list_of_filler_objects:
            all_maslo += (i.mass * i.maslo) / all_mass

        for i in self.list_of_pigmpast_objects:
            all_maslo += (i.mass * i.suhoi_pigm
                          * i.maslo) / all_mass

        self.oil = all_maslo

    def all_const_pigm_in_objects(self):
        all_maslo = self.oil
        degree_pigm = self.degree_pigm
        const_pigm = all_maslo * degree_pigm
        self.const_pigm = const_pigm

    def all_hiding_in_objects(self):
        all_mass = self.all_mass_for_suhoi
        all_mass_pigm = Decimal(0)
        all_hiding = Decimal(0)
        for i in self.list_of_pigment_objects:
            i = i.mass
            all_mass_pigm += i
        for i in self.list_of_pigmpast_objects:
            i = i.mass * i.suhoi_pigm
            all_mass_pigm += i

        for i in self.list_of_pigment_objects:
            all_hiding += (i.mass * i.hiding) / all_mass_pigm

        for i in self.list_of_pigmpast_objects:
            all_hiding += (i.mass * i.suhoi_pigm
                           * i.hiding) / all_mass_pigm

        self.hiding_pigm = all_hiding

        suhoi = self.suhoi
        if all_mass != Decimal(0) and all_mass_pigm != Decimal(0):
            hiding_lkp = (all_hiding * suhoi * 100) / (all_mass_pigm * 100 / all_mass)
        else:
            hiding_lkp = Decimal(0)

        self.hiding_dry = hiding_lkp / 100

        if suhoi != Decimal(0):
            wet_hiding_lkp = hiding_lkp / suhoi
        else:
            wet_hiding_lkp = Decimal(0)

        self.hiding_wet = wet_hiding_lkp

    def filum_in_objects(self):
        all_mass_pigm = Decimal(0)
        filum = Decimal(0)

        try:
            for i in self.list_of_pigment_objects:
                i = i.mass
                all_mass_pigm += i
            for i in self.list_of_pigmpast_objects:
                i = i.mass * i.suhoi_pigm
                all_mass_pigm += i

            for i in self.list_of_pigment_objects:
                filum += (i.mass * i.maslo *
                          i.hiding) / all_mass_pigm
            for i in self.list_of_pigmpast_objects:
                filum += (i.mass * i.suhoi_pigm *
                          i.hiding
                          * i.maslo) / all_mass_pigm
        except Exception as e:
            logging.error(e, exc_info=True)

            filum = Decimal(0)
        filum = filum / Decimal(100)

        self.philum = filum

    def all_kokp_maslo_in_objects(self):
        all_maslo = self.oil
        filler_and_pigm_mass = Decimal(0)
        film_mass = Decimal(0)
        filler_and_pigm_density = Decimal(0)
        film_density = Decimal(0)

        for i in self.list_of_pigment_objects:
            filler_and_pigm_mass += i.mass
        for i in self.list_of_filler_objects:
            filler_and_pigm_mass += i.mass
        for i in self.list_of_pigmpast_objects:
            filler_and_pigm_mass += i.mass * i.suhoi_pigm

        for i in self.list_of_film_objects:
            film_mass += i.mass * i.suhoi
        for i in self.list_of_hardener_objects:
            film_mass += i.mass_for_params * i.suhoi
        for i in self.list_of_pigmpast_objects:
            film_mass += i.mass * i.suhoi_film
        for i in self.list_of_additive_objects:
            if i.type.lower() == 'пластификатор':
                film_mass += i.mass * i.suhoi

        if film_mass <= 0:
            self.okp_kokp = Decimal(0)
            self.kokp = Decimal(0)
            return

        for i in self.list_of_pigment_objects:
            filler_and_pigm_density += ((i.mass * i.density)
                                        / filler_and_pigm_mass)
        for i in self.list_of_filler_objects:
            filler_and_pigm_density += ((i.mass * i.density)
                                        / filler_and_pigm_mass)
        for i in self.list_of_pigmpast_objects:
            filler_and_pigm_density += ((i.mass * i.suhoi_pigm
                                         * i.density_pigm) / filler_and_pigm_mass)

        for i in self.list_of_film_objects:
            film_density += ((i.mass * i.suhoi
                              * i.density_dry) / film_mass)
        for i in self.list_of_hardener_objects:
            film_density += ((i.mass_for_params * i.suhoi
                              * i.density_dry) / film_mass)
        for i in self.list_of_solvent_objects:
            film_density += ((i.mass * i.suhoi
                              * i.density) / film_mass)
        for i in self.list_of_pigmpast_objects:
            film_density += ((i.mass * i.suhoi_film
                              * i.density_dry) / film_mass)
        for i in self.list_of_additive_objects:
            if i.type.lower() == 'пластификатор':
                film_density += ((i.mass * i.suhoi
                                  * i.density) / film_mass)

        try:
            kokp = 100 / (1 + ((all_maslo * filler_and_pigm_density) / (100 * film_density)))
        except Exception as e:
            logging.error(e, exc_info=True)

            kokp = Decimal(0)

        self.kokp = kokp

        try:
            okp_kokp = self.okp / self.kokp
        except Exception as e:
            logging.error(e, exc_info=True)
            okp_kokp = Decimal(0)

        self.okp_kokp = okp_kokp * Decimal(100)

    def count_delta_e_color(self) -> str:
        alfa1 = self.project_color[0:3]
        r1 = float(int(self.project_color[3:5], 16))
        g1 = float(int(self.project_color[5:7], 16))
        b1 = float(int(self.project_color[7:], 16))

        alfa2 = self.recepture_color[0:3]
        r2 = float(int(self.recepture_color[3:5], 16))
        g2 = float(int(self.recepture_color[5:7], 16))
        b2 = float(int(self.recepture_color[7:], 16))

        # прозрачные исключаем
        if alfa1 == "#00" or alfa2 == "#00":
            return ""

        rgb1 = [r1 / 255, b1 / 255, g1 / 255]
        rgb2 = [r2 / 255, b2 / 255, g2 / 255]

        lab1 = color_kit.rgb2lab(rgb1)
        lab2 = color_kit.rgb2lab(rgb2)

        L1 = round(lab1[0], 1)
        a1 = round(lab1[1], 1)
        b1 = round(lab1[2], 1)

        L2 = round(lab2[0], 1)
        a2 = round(lab2[1], 1)
        b2 = round(lab2[2], 1)

        delta = ((L1 - L2) ** 2 + (a1 - a2) ** 2 + (b1 - b2) ** 2) ** 0.5
        delta = str(round(delta, 2)).replace(".", ",")
        return delta


class SearchCombobox(CustomEntry):
    def __init__(self, parent, list_names):
        super(SearchCombobox, self).__init__(parent)
        self.db = DB()
        # self.setEditable(True)
        # self.setDuplicatesEnabled(False)
        # self.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)

        word_set = set(list_names)
        completer = QCompleter(word_set)
        completer.setCaseSensitivity(QtCore.Qt.CaseSensitivity.CaseInsensitive)
        completer.setCompletionMode(QtWidgets.QCompleter.CompletionMode.PopupCompletion)
        completer.setFilterMode(Qt.MatchFlag.MatchContains)
        self.setCompleter(completer)
        self.update()


class ReceptureSettings(QtWidgets.QWidget):

    def __init__(self, parent: ReceptureWindow):
        super(ReceptureSettings, self).__init__()
        set_window_icon(self)
        self.parent_obj = parent
        self.setObjectName("rec_settings")
        self.setStyleSheet("""
        QWidget#rec_settings{
        background: #f9f9f9;
        }
        """)
        self.resize(400, 300)
        self.dry_type = int(get_suhoi_type())
        self.accurate_density = str(self.parent_obj.recepture_data.accurate_density).replace(".", ",")
        self.k_price = str(self.parent_obj.recepture_data.price_K).replace(".", ",")

        self.verticalLayout = QtWidgets.QVBoxLayout(self)
        self.verticalLayout.setObjectName("verticalLayout")

        self.dry_mass_w = QtWidgets.QWidget(parent=self)
        self.horizontalLayout = QtWidgets.QHBoxLayout(self.dry_mass_w)
        self.type_l = QtWidgets.QLabel(parent=self.dry_mass_w)
        self.type_l.setText("Тип расчета м.д.н.в.")
        self.horizontalLayout.addWidget(self.type_l)

        self.type_dry_mass_count = CustomCombobox(self.dry_mass_w)
        self.type_dry_mass_count.addItem("Все компоненты", userData=1)
        self.type_dry_mass_count.addItem("Ограниченный 1", userData=2)
        self.type_dry_mass_count.addItem("Ограниченный 2", userData=3)
        self.type_dry_mass_count.setCurrentIndex(self.dry_type - 1)
        self.type_dry_mass_count.currentIndexChanged.connect(self.show_description)
        self.horizontalLayout.addWidget(self.type_dry_mass_count)
        spacerItem = QtWidgets.QSpacerItem(20, 10, QtWidgets.QSizePolicy.Policy.Expanding,
                                           QtWidgets.QSizePolicy.Policy.Minimum)
        self.horizontalLayout.addItem(spacerItem)
        self.verticalLayout.addWidget(self.dry_mass_w)

        self.descript_type_l = QtWidgets.QLabel(parent=self)
        self.descript_type_l.setText("Опиание")
        self.descript_type_l.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.descript_type_l.setWordWrap(False)
        self.verticalLayout.addWidget(self.descript_type_l)

        self.price_k_w = QtWidgets.QWidget(parent=self)
        self.horizontalLayout_2 = QtWidgets.QHBoxLayout(self.price_k_w)
        self.price_k_l = QtWidgets.QLabel(parent=self.price_k_w)
        self.price_k_l.setText("Коэффициент стоимости")
        tooltip = "Коэффициент, на который будет домножаться стоимость"
        self.price_k_l.setToolTip(tooltip)
        self.horizontalLayout_2.addWidget(self.price_k_l)
        self.price_k_e = CustomEntry(self.price_k_w, padding=False)
        self.price_k_e.setToolTip(tooltip)
        self.price_k_e.setMaximumSize(50, 999)
        self.price_k_e.setValidator(get_numeric_validator())
        self.price_k_e.setText(self.k_price)
        self.horizontalLayout_2.addWidget(self.price_k_e)
        spacerItem = QtWidgets.QSpacerItem(20, 10, QtWidgets.QSizePolicy.Policy.Expanding,
                                           QtWidgets.QSizePolicy.Policy.Minimum)
        self.horizontalLayout_2.addItem(spacerItem)
        self.verticalLayout.addWidget(self.price_k_w)

        self.density_w = QtWidgets.QWidget(parent=self)
        self.horizontalLayout_3 = QtWidgets.QHBoxLayout(self.density_w)
        self.density_l = QtWidgets.QLabel(parent=self.density_w)
        self.density_l.setText("Экспериментальная плотность, г/мл³")
        tooltip = "Плотность, полученная экспериментально. \nНеобходимо для более точного \nрасчета."
        self.density_l.setToolTip(tooltip)
        self.horizontalLayout_3.addWidget(self.density_l)
        self.density_e = CustomEntry(self.density_w, padding=False)
        self.density_e.setMaximumSize(50, 999)
        self.density_e.setToolTip(tooltip)
        self.density_e.setValidator(get_numeric_validator())
        self.density_e.setText(self.accurate_density)
        self.horizontalLayout_3.addWidget(self.density_e)
        spacerItem = QtWidgets.QSpacerItem(20, 10, QtWidgets.QSizePolicy.Policy.Expanding,
                                           QtWidgets.QSizePolicy.Policy.Minimum)
        self.horizontalLayout_3.addItem(spacerItem)
        self.verticalLayout.addWidget(self.density_w)

        spacerItem = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                           QtWidgets.QSizePolicy.Policy.Expanding)
        self.verticalLayout.addItem(spacerItem)
        self.save_btn = DarkBtn_Ui(self, "save_settings")
        self.save_btn.clicked.connect(lambda : self.save())
        self.verticalLayout.addWidget(self.save_btn)

        self.show_description()

        self.setWindowTitle("Настройки расчета рецептур")

    def closeEvent(self, event):
        self.parent_obj.settings_window = None

    def show_description(self, event=None):
        if self.type_dry_mass_count.currentData() == 1:
            self.descript_type_l.setText('При расчете учитываются все категории компонентов')
        if self.type_dry_mass_count.currentData() == 2:
            self.descript_type_l.setText('При расчете не учитываются функциональные добавки,\nкроме пластификаторов')
        if self.type_dry_mass_count.currentData() == 3:
            self.descript_type_l.setText(
                'При расчете не учитываются функциональные добавки \nи отвердители, кроме пластификаторов')

    def save(self):
        price_k = float(self.price_k_e.text().replace(",","."))
        self.parent_obj.recepture_data.price_K = price_k
        density = float(self.density_e.text().replace(",","."))
        self.parent_obj.recepture_data.accurate_density = density

        update_config_param("suhoi_type",str(self.type_dry_mass_count.currentIndex() + 1))
        self.closeEvent(None)
        self.destroy()


class Component:
    db = DB()

    def __init__(self, name, mass, category):
        self.db = Component.db
        self.name = name
        self.mass = Decimal(mass.replace(",", "."))
        valuta = self.db.get_info_reactive(category, self.name, 'valuta')[0][0]
        price = self.get_and_fix_from_db(category, self.name, 'price')
        try:
            if valuta != 'Руб':
                self.price = self.convert_price(valuta, price)
            else:
                self.price = price
        except Exception as e:
            logging.error("Не удалось перевести цену", e)
            self.price = Decimal(0)
        self.dry_mass = None

    def get_and_fix_from_db(self, category, component, param) -> Decimal:
        result = self.db.get_info_reactive(category, component, param)[0][0].replace(",", ".")
        try:
            result = Decimal(result)
        except Exception as e:
            result = Decimal(0)
            logging.error(f"Ошибка при создании объекта {self.name}", e)
        return result

    def convert_price(self, valuta: str, price: Decimal) -> Decimal:
        get_xml = requests.get(
            'http://www.cbr.ru/scripts/XML_daily.asp'
        )
        exchange_rate = {}
        # Парсинг XML используя ElementTree
        structure = ET.fromstring(get_xml.content)

        # Поиск курса доллара (USD ID: R01235)
        dollar = structure.find("./*[@ID='R01235']/Value")
        exchange_rate['$'] = dollar.text.replace(',', '.')

        # Поиск курса евро (EUR ID: R01239)
        euro = structure.find("./*[@ID='R01239']/Value")
        exchange_rate['€'] = euro.text.replace(',', '.')
        converted_price = price * Decimal(exchange_rate[valuta])

        return converted_price


class Solvents(Component):
    def __init__(self, name, mass):
        super(Solvents, self).__init__(name, mass, 'Solvents')
        self.density = self.get_and_fix_from_db('Solvents', self.name, 'density')
        self.density_solvent = self.density
        self.suhoi = Decimal('0')
        self.dry_mass = Decimal('0')
        # print(self.to_string())

    def to_string(self):
        return f"""Растворитель:
        Цена: {self.price}
        Плотность: {self.density}"""


class Pigments(Component):
    def __init__(self, name, mass):
        super(Pigments, self).__init__(name, mass, 'Pigments')
        self.density = self.get_and_fix_from_db('Pigments', self.name, 'density')
        self.maslo = self.get_and_fix_from_db('Pigments', self.name, 'maslo')
        self.hiding = self.get_and_fix_from_db('Pigments', self.name, 'hiding')
        self.suhoi = Decimal("1")
        self.dry_mass = Decimal(mass.replace(",", "."))

    def to_string(self):
        return f"""Пигмент:
        Цена: {self.price}
        Плотность: {self.density}
        Маслоемкость: {self.maslo}
        Укрывистость: {self.hiding}"""


class Fillers(Component):
    def __init__(self, name, mass):
        super(Fillers, self).__init__(name, mass, 'Fillers')
        self.density = self.get_and_fix_from_db('Fillers', self.name, 'density')
        self.maslo = self.get_and_fix_from_db('Fillers', self.name, 'maslo')
        self.suhoi = Decimal("1")
        self.dry_mass = Decimal(mass.replace(",", "."))

    def to_string(self):
        return f"""Наполнитель:
        Цена: {self.price}
        Плотность: {self.density}
        Маслоемкость: {self.maslo}"""


class Films(Component):
    def __init__(self, name, mass):
        super(Films, self).__init__(name, mass, 'Films')
        self.suhoi = self.get_and_fix_from_db('Films', self.name, 'suhoi')
        self.density_dry = self.get_and_fix_from_db('Films', self.name, 'density_dry')
        self.density = self.get_and_fix_from_db('Films', self.name, 'density')
        self.density_solvent = self.get_and_fix_from_db('Films', self.name, 'density_solvent')
        self.dry_mass = Decimal(mass.replace(",", ".")) * self.suhoi
        self.func_groups = self.get_and_fix_from_db('Films', self.name, 'func_groups')

    def to_string(self):
        return f"""Пленкообразователь:
        Цена: {self.price}
        Плотность: {self.density}
        Плотность сухой пленки: {self.density_dry}
        Плотность растворителя: {self.density_solvent}
        Массовая доля нелетучих веществ: {self.suhoi}
        Грамм/эквивалент функциональных групп: {self.func_groups}"""


class Additives(Component):
    def __init__(self, name, mass):
        super(Additives, self).__init__(name, mass, 'Additives')
        self.suhoi = self.get_and_fix_from_db('Additives', self.name, 'suhoi')
        self.dosage = self.db.get_info_reactive('Additives', self.name, 'dosage')[0][0]
        self.density = self.get_and_fix_from_db('Additives', self.name, 'density')
        self.type = self.db.get_info_reactive('Additives', self.name, 'type')[0][0]
        self.density_solvent = self.get_and_fix_from_db('Additives', self.name, 'density_solvent')
        self.dry_mass = Decimal(mass.replace(",", ".")) * self.suhoi

    def to_string(self):
        return f"""Функциональная добавка:
           Цена: {self.price}
           Плотность: {self.density}
           Тип: {self.type}
           Плотность растворителя: {self.density_solvent}
           Массовая доля нелетучих веществ: {self.suhoi}"""


class Pigmpasts(Component):
    def __init__(self, name, mass):
        super(Pigmpasts, self).__init__(name, mass, 'Pigmpast')
        self.suhoi = self.get_and_fix_from_db('Pigmpast', self.name, 'suhoi')
        self.suhoi_pigm = self.get_and_fix_from_db('Pigmpast', self.name, 'suhoi_pigm')
        self.suhoi_film = self.get_and_fix_from_db('Pigmpast', self.name, 'suhoi_film')
        self.maslo = self.get_and_fix_from_db('Pigmpast', self.name, 'maslo')
        self.density = self.get_and_fix_from_db('Pigmpast', self.name, 'density')
        self.density_dry = self.get_and_fix_from_db('Pigmpast', self.name, 'density_dry')
        self.density_pigm = self.get_and_fix_from_db('Pigmpast', self.name, 'density_pigm')
        self.hiding = self.get_and_fix_from_db('Pigmpast', self.name, 'hiding')
        self.density_solvent = self.get_and_fix_from_db('Pigmpast', self.name, 'density_solvent')
        self.dry_mass = Decimal(mass.replace(",", ".")) * self.suhoi

    def to_string(self):
        return f"""Пигментная паста:
        Цена: {self.price}
        Массовая доля нелетучих веществ: {self.suhoi}
        Массовая доля наполнения: {self.suhoi_pigm}
        Массовая доля пленкообразователя: {self.suhoi_film}
        Плотность пасты: {self.density}
        Плотность растворителя: {self.density_solvent}
        Плотность сухой пленки: {self.density_dry}
        Средняя плотность наполнения: {self.density_pigm}
        Средняя маслоемкость наполнения: {self.maslo}
        Укрывистость пигментов: {self.hiding}"""


class Hardeners(Component):
    def __init__(self, name, mass):
        super(Hardeners, self).__init__(name, mass, 'Hardener')
        self.suhoi = self.get_and_fix_from_db('Hardener', self.name, 'suhoi')
        self.func_groups = self.db.get_info_reactive('Hardener', self.name, 'func_groups')[0][0]
        self.density = self.get_and_fix_from_db('Hardener', self.name, 'density')
        self.density_dry = self.get_and_fix_from_db('Hardener', self.name, 'density_dry')
        countable = self.db.get_info_reactive('Hardener', self.name, 'countable')[0][0]
        if countable.lower().strip() == 'да':
            self.mass_for_params = self.mass
        else:
            self.mass_for_params = Decimal('0')  # для расчета окп кокп СП
        self.density_solvent = self.get_and_fix_from_db('Hardener', self.name, 'density_solvent')
        self.dry_mass = Decimal(mass.replace(",", ".")) * self.suhoi

    def to_string(self):
        return f"""Отвердитель:
           Цена: {self.price}
           Плотность: {self.density}
           Плотность сухого компонента: {self.density_dry}
           Плотность растворителя: {self.density_solvent}
           Массовая доля нелетучих веществ: {self.suhoi}"""


class CountAdditiveWindow(QtWidgets.QWidget):
    def __init__(self, parent: ReceptureWindow):
        super(CountAdditiveWindow, self).__init__()
        self.recepture = parent
        self.first_row = 2
        self.list_checkbox = []
        self.list_components = self.recepture.recepture_data.get_components_obj()
        self.additive_obj = None
        self.db = DB()
        self.type = None

        set_window_icon(self)
        self.setObjectName("CountAdditiveWindow")
        self.setStyleSheet("""
        QWidget#CountAdditiveWindow{
        background: #f9f9f9;
        }
        """)

        self.resize(537, 366)
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.list_comp_w = QtWidgets.QWidget(parent=self)
        self.list_comp_w.setObjectName("list_comp_w")
        self.list_comp_w.setStyleSheet("""
                QWidget#list_comp_w{
                border-right: 2px solid #eee;
                }
                """)
        self.gridLayout = QtWidgets.QGridLayout(self.list_comp_w)
        self.gridLayout.setObjectName("gridLayout")

        self.check_l = QtWidgets.QLabel(parent=self.list_comp_w)
        self.gridLayout.addWidget(self.check_l, 1, 0, 1, 1)
        self.mass_l = QtWidgets.QLabel(parent=self.list_comp_w)
        self.gridLayout.addWidget(self.mass_l, 1, 1, 1, 1)

        self.comp_l = QtWidgets.QLabel(parent=self.list_comp_w)
        self.gridLayout.addWidget(self.comp_l, 0, 0, 1, 2)

        self.dry_mass_l = QtWidgets.QLabel(parent=self.list_comp_w)
        self.dry_mass_l.setObjectName("dry_mass_l")
        self.gridLayout.addWidget(self.dry_mass_l, 1, 2, 1, 1)

        for obj in self.list_components:
            self.add_row_component(obj)

        spacerItem = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                           QtWidgets.QSizePolicy.Policy.Expanding)
        self.gridLayout.addItem(spacerItem, self.first_row, 1, 1, 1)
        self.horizontalLayout.addWidget(self.list_comp_w)

# right part:
        self.count_w = QtWidgets.QWidget(parent=self)
        self.gridLayout_2 = QtWidgets.QGridLayout(self.count_w)
        self.gridLayout_2.setVerticalSpacing(2)

        self.list_all_names = self.db.load_reactives("Additives", "name")
        self.list_all_names = list(map(lambda x: x[0], self.list_all_names))
        self.component_e = SearchCombobox(self, self.list_all_names)
        self.component_e.setMinimumSize(QtCore.QSize(250, 0))
        self.component_e.textChanged.connect(self.name_changed)
        self.gridLayout_2.addWidget(self.component_e, 1, 0, 1, 4)

        self.dry_rb = QtWidgets.QRadioButton(parent=self.count_w)
        self.dry_rb.setChecked(True)
        self.gridLayout_2.addWidget(self.dry_rb, 5, 0, 1, 3)
        self.all_mass_rb = QtWidgets.QRadioButton(parent=self.count_w)
        self.gridLayout_2.addWidget(self.all_mass_rb, 6, 0, 1, 3)

        self.metal_w = QtWidgets.QWidget(parent=self.count_w)
        lo = QtWidgets.QHBoxLayout(self.metal_w)
        lo.setContentsMargins(0,0,0,0)

        metal_l = QtWidgets.QLabel(self.metal_w)
        metal_l.setText("Содержание металла, %")
        lo.addWidget(metal_l)
        self.metal_e = CustomEntry(self.metal_w, padding=False)
        self.metal_e.setMaximumSize(50, 22)
        self.metal_e.setValidator(get_numeric_validator())
        lo.addWidget(self.metal_e)
        spacerItem1 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Expanding,
                                            QtWidgets.QSizePolicy.Policy.Minimum)
        lo.addItem(spacerItem1)
        self.gridLayout_2.addWidget(self.metal_w, 6, 0, 1, 4)
        self.metal_w.hide()

        dosage_w = QtWidgets.QWidget(parent=self.count_w)
        lo = QtWidgets.QHBoxLayout(dosage_w)
        lo.setContentsMargins(0, 0, 0, 0)
        dosage_l = QtWidgets.QLabel(parent=dosage_w)
        dosage_l.setText("Дозировка, %")
        lo.addWidget(dosage_l)
        self.dosage_e = CustomEntry(dosage_w, padding=False)
        self.dosage_e.setMaximumSize(50, 22)
        self.dosage_e.setValidator(get_numeric_validator())
        lo.addWidget(self.dosage_e)
        count_new_b = ColorButton(parent=dosage_w, color="blue")
        count_new_b.clicked.connect(lambda: self.count())
        count_new_b.setText("Рассчитать")
        lo.addWidget(count_new_b)
        spacerItem1 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Expanding,
                                            QtWidgets.QSizePolicy.Policy.Minimum)
        lo.addItem(spacerItem1)
        self.gridLayout_2.addWidget(dosage_w, 7, 0, 1, 4)

        self.info = QtWidgets.QLabel(parent=self.count_w)
        self.gridLayout_2.addWidget(self.info, 4, 0, 1, 5)

        result_w = QtWidgets.QWidget(parent=self.count_w)
        lo = QtWidgets.QHBoxLayout(result_w)
        lo.setContentsMargins(0, 0, 0, 0)
        result_l = QtWidgets.QLabel(parent=result_w)
        result_l.setText("Результат:")
        lo.addWidget(result_l)
        self.result_e = CustomEntry(self.count_w, padding=False)
        self.result_e.setMaximumSize(50, 22)
        self.result_e.setReadOnly(True)
        lo.addWidget(self.result_e)
        add_count_b = ColorButton(parent=result_w, color="blue")
        add_count_b.setText("Очистить")
        add_count_b.clicked.connect(lambda: self.clear())
        lo.addWidget(add_count_b)
        spacerItem1 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Expanding,
                                            QtWidgets.QSizePolicy.Policy.Minimum)
        lo.addItem(spacerItem1)
        self.gridLayout_2.addWidget(result_w, 9, 0, 1, 1)

        self.additive_name_l = QtWidgets.QLabel(parent=self.count_w)
        self.gridLayout_2.addWidget(self.additive_name_l, 0, 0, 1, 4)
        self.info_l = QtWidgets.QLabel(parent=self.count_w)
        self.gridLayout_2.addWidget(self.info_l, 2, 0, 1, 3)

        spacerItem1 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                            QtWidgets.QSizePolicy.Policy.Expanding)
        self.gridLayout_2.addItem(spacerItem1, 4, 0, 1, 1)
        self.horizontalLayout.addWidget(self.count_w)

        self.setWindowTitle("Расчет функциональных добавок")

        self.check_l.setText("Учитывать")
        self.mass_l.setText("Масса")
        self.comp_l.setText("Компоненты")

        self.comp_l.setFont(generate_font(12))

        self.dry_mass_l.setText("м.н.в.")
        self.dry_rb.setText("По массе нелетучих веществ")
        self.additive_name_l.setText("Функциональная добавка:")
        self.additive_name_l.setFont(generate_font(12))
        self.info_l.setText("Информация о дозировке:")
        self.all_mass_rb.setText("По всей массе")

    def closeEvent(self, event):
        self.recepture.additive_window = None

    def add_row_component(self, comp_obj: Component):
        component_chek = QtWidgets.QCheckBox(parent=self.list_comp_w)
        component_chek.user_data = comp_obj
        self.list_checkbox.append(component_chek)
        component_chek.setText(comp_obj.name)

        self.gridLayout.addWidget(component_chek, self.first_row, 0, 1, 1)

        mass = QtWidgets.QLabel(parent=self.list_comp_w)
        mass.setText(normalize_number(comp_obj.mass))
        self.gridLayout.addWidget(mass, self.first_row, 1, 1, 1)

        dry_mass = QtWidgets.QLabel(parent=self.list_comp_w)
        dry_mass.setText(normalize_number(comp_obj.dry_mass))
        self.gridLayout.addWidget(dry_mass, self.first_row, 2, 1, 1)

        self.first_row += 1

    def name_changed(self, text):
        if text in self.list_all_names:
            self.additive_obj = Additives(text, "0")
            info = self.additive_obj.dosage

            if self.additive_obj.type.lower() == "сиккатив":
                self.type = "Cиккатив"
                self.dry_rb.hide()
                self.all_mass_rb.hide()
                self.metal_w.show()
                info += '\n\nСправочная дозировка для сухого алкида: \nКобальт - 0,03-0,05%' \
                              ' \tМарганец - 0,02-0,04% \tСвинец - 0,06-0,08%' \
                              ' \nЦирконий - 0,08-0,15% \tКальций - 0,05-2,00% ' \
                              ' \tВанадий - 0,02-0,07% \nЖелезо - 0,02-0,05% \tЦерий - 0,1-0,2% \tКальций - 0,05-2,0% ' \
                              '\nАлюминий - 0,2-1,0% \tЦинк - 0,05-0,25% \tВисмут - 0,02-0,1% ' \
                              '\nСтронций - 0,1-0,5% \tБарий - 0,1-0,25% \tЛитий - 0,1-0,02% ' \
                              '\nРедкоземельные(La, Nd) - 0,1-0,3%' \
                              '\nДля смесевых сиккативов расчет ведется по ' \
                              'основному металлу \n(дающему наиболее желаемые свойства).' \
                              ' \nАнтипленка,например, МЕКО добавляется в количестве 10% от массы сиккатива.'
            else:
                self.dry_rb.show()
                self.all_mass_rb.show()
                self.metal_w.hide()
                self.type = "Добавка"

            self.info.setText(info)
        else:
            self.additive_obj = None
            self.type = None
            self.info.setText("")
            self.clear()

    def count(self):
        if self.type == "Добавка":
            self.count_additive()
        elif self.type == "Cиккатив":
            self.count_sickative()
        else:
            # print(self.type)
            InfoWindow("Укажите название имеющейся функциональной добавки").exec()


    def collect_mass(self) -> Decimal:
        mass = Decimal(0)
        checkbox: QtWidgets.QCheckBox
        for checkbox in self.list_checkbox:
            if checkbox.isChecked():
                comp_obj: Component = checkbox.user_data
                if self.dry_rb.isChecked() or self.type == "Сиккатив":
                    mass += comp_obj.dry_mass
                else:
                    mass += comp_obj.mass
        return mass

    def count_sickative(self):
        mass = self.collect_mass()
        metal = self.metal_e.text().replace(",", ".")
        if metal.strip() in ["", "."]:
            metal = "0"
        metal = Decimal(metal)

        dosage = self.dosage_e.text().replace(",", ".")
        if dosage.strip() in ["", "."]:
            dosage = "0"
        dosage = Decimal(dosage)

        result = mass * dosage / metal

        old = self.result_e.text().replace(",", ".")
        if old.strip() in ["", "."]:
            old = "0"
        old = Decimal(old)
        result += old

        result = normalize_number(result)
        self.result_e.setText(result)

    def count_additive(self):
        mass = self.collect_mass()
        dosage = self.dosage_e.text().replace(",", ".")
        if dosage.strip() in ["", "."]:
            dosage = "0"
        dosage = Decimal(dosage)
        result = (mass * dosage) / Decimal(100)
        old = self.result_e.text().replace(",", ".")
        if old.strip() in ["", "."]:
            old = "0"
        old = Decimal(old)
        result += old

        result = normalize_number(result)
        self.result_e.setText(result)

    def clear(self):
        self.result_e.setText("")


class CountHardenerWindow(QtWidgets.QWidget):
    def __init__(self, parent: ReceptureWindow):
        super(CountHardenerWindow, self).__init__()
        set_window_icon(self)

        self.recepture = parent
        self.first_row = 2
        self.list_checkbox = []

        list_components = self.recepture.recepture_data.get_components_obj()
        self.list_components = list(filter(lambda x:  isinstance(x, Films), list_components))
        self.additive_obj = None
        self.db = DB()
        self.type = None

        self.setObjectName("CountAdditiveWindow")
        self.setStyleSheet("""
        QWidget#CountAdditiveWindow{
        background: #f9f9f9;
        }
        """)

        self.resize(537, 366)
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.list_comp_w = QtWidgets.QWidget(parent=self)
        self.list_comp_w.setObjectName("list_comp_w")
        self.list_comp_w.setStyleSheet("""
                QWidget#list_comp_w{
                border-right: 2px solid #eee;
                }
                """)
        self.gridLayout = QtWidgets.QGridLayout(self.list_comp_w)
        self.gridLayout.setObjectName("gridLayout")

        self.check_l = QtWidgets.QLabel(parent=self.list_comp_w)
        self.gridLayout.addWidget(self.check_l, 1, 0, 1, 1)
        self.mass_l = QtWidgets.QLabel(parent=self.list_comp_w)
        self.gridLayout.addWidget(self.mass_l, 1, 1, 1, 1)

        self.comp_l = QtWidgets.QLabel(parent=self.list_comp_w)
        self.gridLayout.addWidget(self.comp_l, 0, 0, 1, 2)

        self.dry_mass_l = QtWidgets.QLabel(parent=self.list_comp_w)
        self.dry_mass_l.setObjectName("dry_mass_l")
        self.gridLayout.addWidget(self.dry_mass_l, 1, 2, 1, 1)

        for obj in self.list_components:
            self.add_row_component(obj)

        spacerItem = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                           QtWidgets.QSizePolicy.Policy.Expanding)
        self.gridLayout.addItem(spacerItem, self.first_row, 1, 1, 1)
        self.horizontalLayout.addWidget(self.list_comp_w)

# right part:
        self.count_w = QtWidgets.QWidget(parent=self)
        self.gridLayout_2 = QtWidgets.QGridLayout(self.count_w)
        self.gridLayout_2.setVerticalSpacing(2)

        self.list_all_names = self.db.load_reactives("Hardener", "name")
        self.list_all_names = list(map(lambda x: x[0], self.list_all_names))
        self.component_e = SearchCombobox(self, self.list_all_names)
        self.component_e.setMinimumSize(QtCore.QSize(250, 0))
        self.component_e.textChanged.connect(self.name_changed)
        self.gridLayout_2.addWidget(self.component_e, 1, 0, 1, 4)

        self.info = QtWidgets.QLabel(parent=self.count_w)
        self.gridLayout_2.addWidget(self.info, 4, 0, 1, 5)

        self.const_w = QtWidgets.QWidget(parent=self.count_w)
        tooltip = """Значения k: 
        \nk=1-1,4 - полиамины
        \nk=0,8-0,9 - ароматические изоцианаты
        \nk=1,05-1,15 - алифатические изоцианаты
        \nk=1 - смесевые изоцианаты
        """
        lo = QtWidgets.QHBoxLayout(self.const_w)
        lo.setContentsMargins(0,0,0,0)
        const_l = QtWidgets.QLabel(self.const_w)
        const_l.setText("Поправочный коэффициент k:")
        const_l.setToolTip(tooltip)
        lo.addWidget(const_l)
        self.const_e = CustomEntry(self.const_w, padding=False)
        self.const_e.setMaximumSize(50, 22)
        self.const_e.setValidator(get_numeric_validator())
        self.const_e.setToolTip(tooltip)
        lo.addWidget(self.const_e)
        spacerItem1 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Expanding,
                                            QtWidgets.QSizePolicy.Policy.Minimum)
        lo.addItem(spacerItem1)
        self.gridLayout_2.addWidget(self.const_w, 6, 0, 1, 4)


        equivalent_w = QtWidgets.QWidget(parent=self.count_w)
        lo = QtWidgets.QHBoxLayout(equivalent_w)
        lo.setContentsMargins(0, 0, 0, 0)
        equivalent_l = QtWidgets.QLabel(parent=equivalent_w)
        equivalent_l.setText("Эквивалентная масса, г/экв:")
        lo.addWidget(equivalent_l)
        self.equivalent_e = CustomEntry(equivalent_w, padding=False)
        self.equivalent_e.setMaximumSize(50, 22)
        self.equivalent_e.setValidator(get_numeric_validator())
        lo.addWidget(self.equivalent_e)

        count_new_b = ColorButton(parent=equivalent_w, color="blue")
        count_new_b.clicked.connect(lambda: self.count())
        count_new_b.setText("Рассчитать")
        lo.addWidget(count_new_b)
        spacerItem1 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Expanding,
                                            QtWidgets.QSizePolicy.Policy.Minimum)
        lo.addItem(spacerItem1)
        self.gridLayout_2.addWidget(equivalent_w, 7, 0, 1, 4)

        result_w = QtWidgets.QWidget(parent=self.count_w)
        lo = QtWidgets.QHBoxLayout(result_w)
        lo.setContentsMargins(0, 0, 0, 0)
        result_l = QtWidgets.QLabel(parent=result_w)
        result_l.setText("Результат:")
        lo.addWidget(result_l)
        self.result_e = CustomEntry(self.count_w, padding=False)
        self.result_e.setMaximumSize(50, 22)
        self.result_e.setReadOnly(True)
        lo.addWidget(self.result_e)
        add_count_b = ColorButton(parent=result_w, color="blue")
        add_count_b.setText("Очистить")
        add_count_b.clicked.connect(lambda: self.clear())
        lo.addWidget(add_count_b)
        spacerItem1 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Expanding,
                                            QtWidgets.QSizePolicy.Policy.Minimum)
        lo.addItem(spacerItem1)
        self.gridLayout_2.addWidget(result_w, 9, 0, 1, 1)

        self.additive_name_l = QtWidgets.QLabel(parent=self.count_w)
        self.gridLayout_2.addWidget(self.additive_name_l, 0, 0, 1, 4)
        self.info_l = QtWidgets.QLabel(parent=self.count_w)
        self.gridLayout_2.addWidget(self.info_l, 2, 0, 1, 3)

        spacerItem1 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                            QtWidgets.QSizePolicy.Policy.Expanding)
        self.gridLayout_2.addItem(spacerItem1, 4, 0, 1, 1)
        self.horizontalLayout.addWidget(self.count_w)

        self.setWindowTitle("Расчет отвердителей")

        self.check_l.setText("Учитывать")
        self.mass_l.setText("Масса")
        self.comp_l.setText("Пленкообразователи")

        self.comp_l.setFont(generate_font(12))

        self.dry_mass_l.setText("м.н.в.")

        self.additive_name_l.setText("Отвердитель:")
        self.additive_name_l.setFont(generate_font(12))
        self.info_l.setText("Информация об эквивалентных массах:")

    def closeEvent(self, event):
        self.recepture.hardener_window = None

    def add_row_component(self, comp_obj: Component):
        component_chek = QtWidgets.QCheckBox(parent=self.list_comp_w)
        component_chek.user_data = comp_obj
        self.list_checkbox.append(component_chek)
        component_chek.setText(comp_obj.name)

        self.gridLayout.addWidget(component_chek, self.first_row, 0, 1, 1)

        mass = QtWidgets.QLabel(parent=self.list_comp_w)
        mass.setText(normalize_number(comp_obj.mass))
        self.gridLayout.addWidget(mass, self.first_row, 1, 1, 1)

        dry_mass = QtWidgets.QLabel(parent=self.list_comp_w)
        dry_mass.setText(normalize_number(comp_obj.dry_mass))
        self.gridLayout.addWidget(dry_mass, self.first_row, 2, 1, 1)

        self.first_row += 1

    def name_changed(self, text):
        if text in self.list_all_names:
            self.hardener_obj = Hardeners(text, "0")
            info = self.hardener_obj.func_groups
            self.info.setText(info)
        else:
            self.hardener_obj = None
            self.info.setText("")
            self.clear()

    def count(self):
        if self.hardener_obj is not None:
            self.count_hardener()
        else:
            InfoWindow("Укажите название имеющегося отвердителя").exec()

    def collect_film_obj(self) -> List:
        checkbox: QtWidgets.QCheckBox
        list_obj = []
        for checkbox in self.list_checkbox:
            if checkbox.isChecked():
                comp_obj: Films = checkbox.user_data
                list_obj.append(comp_obj)
        return list_obj

    def count_hardener(self):
        list_film_obj = self.collect_film_obj()
        result = Decimal("0")

        equivalent = self.equivalent_e.text().replace(",", ".")
        if equivalent.strip() in ["", "."]:
            equivalent = "0"
        equivalent = Decimal(equivalent)

        const = self.const_e.text().replace(",", ".")
        if const.strip() in ["", "."]:
            const = "0"
        const = Decimal(const)
        error_components = ""
        for film in list_film_obj:
            film: Films
            try:
                result += const * film.dry_mass * equivalent / (film.func_groups * self.hardener_obj.suhoi)
            except Exception:
                error_components += f"{film.name}\n"

        if len(error_components) > 1:
            InfoWindow("Ошибка в расчетах! \nПроверьте в базе наличие функциональных групп \nу следующих компонентов:\n"
                       f"{error_components}"
                       f"И м.д.н.в. у {self.hardener_obj.name}").exec()

        old = self.result_e.text().replace(",", ".")
        if old.strip() in ["", "."]:
            old = "0"
        old = Decimal(old)
        result += old

        result = normalize_number(result)
        self.result_e.setText(result)

    def clear(self):
        self.result_e.setText("")


class RecountOnMaslo(QtWidgets.QWidget):
    def __init__(self, parent: ReceptureWindow):
        super(RecountOnMaslo, self).__init__()
        set_window_icon(self)
        self.recepture = parent
        self.first_row = 2
        self.list_checkbox = []

        list_components = self.recepture.recepture_data.get_components_obj()
        self.list_components = list(filter(lambda x:  isinstance(x, (Pigments, Fillers)), list_components))
        self.additive_obj = None
        self.db = DB()
        self.type = None

        self.setObjectName("CountAdditiveWindow")
        self.setStyleSheet("""
        QWidget#CountAdditiveWindow{
        background: #f9f9f9;
        }
        """)

        self.resize(537, 366)
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.list_comp_w = QtWidgets.QWidget(parent=self)
        self.list_comp_w.setObjectName("list_comp_w")
        self.list_comp_w.setStyleSheet("""
        QWidget#list_comp_w{
        border-right: 2px solid #eee;
        }
        """)

        self.gridLayout = QtWidgets.QGridLayout(self.list_comp_w)
        self.gridLayout.setObjectName("gridLayout")

        self.check_l = QtWidgets.QLabel(parent=self.list_comp_w)
        self.gridLayout.addWidget(self.check_l, 1, 0, 1, 1)
        self.mass_l = QtWidgets.QLabel(parent=self.list_comp_w)
        self.gridLayout.addWidget(self.mass_l, 1, 1, 1, 1)

        self.comp_l = QtWidgets.QLabel(parent=self.list_comp_w)
        self.gridLayout.addWidget(self.comp_l, 0, 0, 1, 2)

        self.dry_mass_l = QtWidgets.QLabel(parent=self.list_comp_w)
        self.dry_mass_l.setObjectName("dry_mass_l")
        self.gridLayout.addWidget(self.dry_mass_l, 1, 2, 1, 1)

        for obj in self.list_components:
            self.add_row_component(obj)

        spacerItem = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                           QtWidgets.QSizePolicy.Policy.Expanding)
        self.gridLayout.addItem(spacerItem, self.first_row, 1, 1, 1)
        self.horizontalLayout.addWidget(self.list_comp_w)

# right part:
        self.count_w = QtWidgets.QWidget(parent=self)
        self.gridLayout_2 = QtWidgets.QGridLayout(self.count_w)
        self.gridLayout_2.setVerticalSpacing(2)

        self.list_all_pigments = self.db.load_reactives("Pigments", "name")
        self.list_all_pigments = list(map(lambda x: x[0], self.list_all_pigments))
        self.list_all_fillers = self.db.load_reactives("Fillers", "name")
        self.list_all_fillers = list(map(lambda x: x[0], self.list_all_fillers))
        list_all_names = self.list_all_pigments + self.list_all_fillers
        self.component_e = SearchCombobox(self, list_all_names)
        self.component_e.setMinimumSize(QtCore.QSize(250, 0))
        self.component_e.textChanged.connect(self.name_changed)
        self.gridLayout_2.addWidget(self.component_e, 1, 0, 1, 4)

        result_w = QtWidgets.QWidget(parent=self.count_w)
        lo = QtWidgets.QHBoxLayout(result_w)
        lo.setContentsMargins(0, 0, 0, 0)
        result_l = QtWidgets.QLabel(parent=result_w)
        result_l.setText("Результат:")
        lo.addWidget(result_l)
        self.result_e = CustomEntry(self.count_w, padding=False)
        self.result_e.setMaximumSize(50, 22)
        self.result_e.setReadOnly(True)
        lo.addWidget(self.result_e)
        count_new_b = ColorButton(parent=result_w, color="blue")
        count_new_b.clicked.connect(lambda: self.count())
        count_new_b.setText("Рассчитать")
        lo.addWidget(count_new_b)
        spacerItem1 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Expanding,
                                            QtWidgets.QSizePolicy.Policy.Minimum)
        lo.addItem(spacerItem1)
        self.gridLayout_2.addWidget(result_w, 9, 0, 1, 1)

        self.additive_name_l = QtWidgets.QLabel(parent=self.count_w)
        self.gridLayout_2.addWidget(self.additive_name_l, 0, 0, 1, 4)
        self.info_l = QtWidgets.QLabel(parent=self.count_w)
        self.gridLayout_2.addWidget(self.info_l, 2, 0, 1, 3)

        spacerItem1 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                            QtWidgets.QSizePolicy.Policy.Expanding)
        self.gridLayout_2.addItem(spacerItem1, 4, 0, 1, 1)
        self.horizontalLayout.addWidget(self.count_w)

        self.setWindowTitle("Замена по маслоемкости")

        self.check_l.setText("Учитывать")
        self.mass_l.setText("Масса")
        self.comp_l.setText("Компоненты")

        self.comp_l.setFont(generate_font(12))

        self.additive_name_l.setText("Заменить на:")
        self.additive_name_l.setFont(generate_font(12))

    def closeEvent(self, event):
        self.recepture.recount_on_maslo_window = None

    def add_row_component(self, comp_obj: Component):
        component_chek = QtWidgets.QCheckBox(parent=self.list_comp_w)
        component_chek.user_data = comp_obj
        self.list_checkbox.append(component_chek)
        component_chek.setText(comp_obj.name)

        self.gridLayout.addWidget(component_chek, self.first_row, 0, 1, 1)

        mass = QtWidgets.QLabel(parent=self.list_comp_w)
        mass.setText(normalize_number(comp_obj.mass))
        self.gridLayout.addWidget(mass, self.first_row, 1, 1, 1)

        self.first_row += 1

    def name_changed(self, text):
        if text in self.list_all_pigments:
            self.new_comp_obj = Pigments(text, "0")
        elif text in self.list_all_fillers:
            self.new_comp_obj = Fillers(text, "0")
        else:
            self.new_comp_obj = None
            self.clear()

    def count(self):
        if self.new_comp_obj is not None:
            self.recount_on_maslo()
        else:
            InfoWindow("Укажите название имеющегося компонента").exec()

    def collect_comp_obj(self) -> List:
        checkbox: QtWidgets.QCheckBox
        list_obj = []
        for checkbox in self.list_checkbox:
            if checkbox.isChecked():
                comp_obj: Fillers or Pigments = checkbox.user_data
                list_obj.append(comp_obj)
        return list_obj

    def recount_on_maslo(self):
        list_comp_obj = self.collect_comp_obj()
        result = Decimal("0")

        error_f = False
        for component in list_comp_obj:
            component: Fillers or Pigments
            try:
                result += component.mass * component.maslo / self.new_comp_obj.maslo
            except Exception:
                error_f = True

        if error_f:
            InfoWindow("Ошибка в расчетах! \nПроверьте в базе правильность маслоемкости \n"
                       f"у{self.new_comp_obj.name}").exec()

        result = normalize_number(result)
        self.result_e.setText(result)

    def clear(self):
        self.result_e.setText("")


class CountReceptureConstant(QtWidgets.QWidget):
    def __init__(self, parent: ReceptureWindow):
        super(CountReceptureConstant, self).__init__()
        set_window_icon(self)
        self.setWindowTitle("Пересчет рецептуры по константе наполнения")
        self.recepture = parent
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.horizontalLayout.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout.setSpacing(0)
        self.setObjectName("CountRecepture")
        self.setStyleSheet("""
        QWidget#CountRecepture{
        background: #f9f9f9;
        }
        """)

        self.list_curve_w = QtWidgets.QWidget(parent=self)
        self.list_curve_w.setMaximumSize(QtCore.QSize(300, 16777215))
        self.list_curve_w.setMinimumSize(QtCore.QSize(250, 16777215))
        self.list_curve_w.setObjectName("ListCurveW")
        self.list_curve_w.setStyleSheet("""
        QWidget#ListCurveW{
        background: white;
         border-right: 2px solid #eee;
        }
        """)

        self.verticalLayout = QtWidgets.QVBoxLayout(self.list_curve_w)
        self.verticalLayout.setContentsMargins(0, 5, 2, 0)
        self.verticalLayout.setSpacing(2)
        label = QtWidgets.QLabel(self.list_curve_w)
        label.setText("Тарировочные кривые")
        label.setFont(generate_font(12, bold=True))
        self.verticalLayout.addWidget(label, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)

        self.list_curve_ui = CustomListItem(parent=self.list_curve_w)
        self.list_curve = os.listdir('saves/' + 'Тарировочные_кривые')
        self.list_curve.remove('params')
        self.list_curve_ui.clicked.connect(lambda x: self.select_curve(self.list_curve[x.row()]))
        self.list_curve_ui.set_list_elements(self.list_curve)
        self.verticalLayout.addWidget(self.list_curve_ui)
        self.horizontalLayout.addWidget(self.list_curve_w)

        self.plot_w = QtWidgets.QWidget(parent=self)
        self.verticalLayout_2 = QtWidgets.QVBoxLayout(self.plot_w)
        label = QtWidgets.QLabel(self.plot_w)
        label.setText("Выберите степень пигментирования")
        label.setFont(generate_font(12, bold=True))
        self.verticalLayout_2.addWidget(label, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)
        self.plot = MplCanvas(self.plot_w)
        self.plot.mpl_connect('button_press_event', self.onclick)
        self.verticalLayout_2.addWidget(self.plot)
        self.verticalLayout_2.addItem(get_v_spacer())
        self.horizontalLayout.addWidget(self.plot_w)

        self.result_w = QtWidgets.QWidget(parent=self)
        self.result_w.setObjectName("ResultW")
        self.result_w.setStyleSheet("""
        QWidget#ResultW{
        border-left: 2px solid #eee;
        }
        """)
        self.result_w.setMinimumSize(QtCore.QSize(200, 16777215))
        self.verticalLayout_3 = QtWidgets.QVBoxLayout(self.result_w)
        label = QtWidgets.QLabel(self.result_w)
        label.setText("Результат")
        label.setFont(generate_font(12, bold=True))
        self.verticalLayout_3.addWidget(label, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)
        self.verticalLayout_3.addItem(get_v_spacer())

        self.horizontalLayout.addWidget(self.result_w)

        change_position_window(self, x=-100)

    def closeEvent(self, event):
        self.recepture.count_recepture_constant = None

    def select_curve(self, name: str):
        self.drow_graf(name)

    def drow_graf(self, name_curve):
        try:
            with SqliteDict('saves/' + 'Тарировочные_кривые' + '/' + name_curve) as mydict:
                receptures = list(mydict.keys())

                list_of_glass = []
                list_of_degree = []
                list_of_maslo = []
                list_of_hiding_pigm = []

                for recepture in receptures:
                    list_of_glass.append(mydict[recepture][3][0])
                    list_of_degree.append(mydict[recepture][6][2])
                    list_of_maslo.append(mydict[recepture][6][4])
                    list_of_hiding_pigm.append(mydict[recepture][6][6])

                dec_list_of_glass = list_of_glass
                list_of_glass = list(map(lambda x: float(x.replace(',', '.')), dec_list_of_glass))
                dec_list_of_degree = list_of_degree
                list_of_degree = list(map(lambda x: float(x.replace(',', '.')), dec_list_of_degree))
                dec_list_of_maslo = list_of_maslo

                # print(list_of_glass)
                # print(list_of_degree)

                sorted_x_y = sorted(list(zip(list_of_degree, list_of_glass)), key=lambda x: x[0])
                list_of_degree, list_of_glass = zip(*sorted_x_y)

                maslo = Counter(dec_list_of_maslo).most_common(1)
                self.tarir_maslo = Decimal(maslo[0][0].replace(',', '.'))

                hiding_pigm = Counter(list_of_hiding_pigm).most_common(1)
                self.tarir_hiding_pigm = Decimal(hiding_pigm[0][0].replace(',', '.'))

                # print(maslo)
            self.plot.axes.clear()
            self.plot.axes.plot(list_of_degree, list_of_glass, '*', label=name_curve)
            popt = self.plot.curve_fit(self.plot.exponenta, list_of_degree, list_of_glass)
            self.plot.axes.plot(list_of_degree, self.plot.exponenta(np.array(list_of_degree), *popt), 'r-')

            self.plot.axes.set(xlabel='Степень пигментирования')
            self.plot.axes.set(ylabel='Блеск')

            self.plot.axes.legend(loc='upper right', frameon=False)
            self.plot.axes.grid(linestyle='--')
            self.plot.draw()

        except Exception as e:
            logging.error(e, exc_info=True)
            InfoWindow('Для построения кривой необходимо минимум 5 точек.').exec()

# TODO: Когда-нибудь можно будет переписать это нормально
    def onclick(self, event):
        tarir_degree_pigm = Decimal(float(event.xdata)).quantize(Decimal("1.00"), "ROUND_HALF_EVEN")
        # print(f'tarir_degree_pigm = {tarir_degree_pigm}')
        tarir_const_nap = tarir_degree_pigm * self.tarir_maslo
        # print(f'tarir_const_nap = {tarir_const_nap}')

        current_maslo = self.recepture.recepture_data.oil
        # print(f'current_maslo = {current_maslo}')
        new_degree_pigm = tarir_const_nap / current_maslo
        # print(f'new_degree_pigm = {new_degree_pigm}')

        suhoi = self.recepture.recepture_data.suhoi

        mass_pigms = (new_degree_pigm * suhoi) / (Decimal(1) + new_degree_pigm)
        mass_films = suhoi - mass_pigms  # масса 100% ПО
        list_new_reactives = []
        list_new_mass = []
        list_new_category = []
        list_add_mass_film = []
        list_films = []

        # начало расчета пигментов
        old_all_mass_pigms = Decimal(0)
        for i in self.recepture.recepture_data.list_of_pigment_objects:
            old_all_mass_pigms += i.mass
        for i in self.recepture.recepture_data.list_of_filler_objects:
            old_all_mass_pigms += i.mass
        for i in self.recepture.recepture_data.list_of_pigmpast_objects:
            old_all_mass_pigms += i.mass * i.suhoi_pigm

        for i in self.recepture.recepture_data.list_of_pigment_objects:
            list_new_reactives.append(i.name)
            list_new_category.append('Pigments')
            list_new_mass.append(i.mass * mass_pigms / old_all_mass_pigms)
        for i in self.recepture.recepture_data.list_of_filler_objects:
            list_new_reactives.append(i.name)
            list_new_category.append('Fillers')
            list_new_mass.append(i.mass * mass_pigms / old_all_mass_pigms)
        for i in self.recepture.recepture_data.list_of_pigmpast_objects:
            list_new_reactives.append(i.name)
            list_new_category.append('PigmPast')
            list_new_mass.append(i.mass * mass_pigms / old_all_mass_pigms)
        # конец расчета пигментов

        # начало расчета пленок
        for i in self.recepture.recepture_data.list_of_pigmpast_objects:  # вычитаем то что уже в пигм пастах
            mass_films -= i.mass * i.suhoi_film

        old_mass_films = Decimal(0)

        for i in self.recepture.recepture_data.list_of_film_objects:
            old_mass_films += i.mass
        for i in self.recepture.recepture_data.list_of_additive_objects:
            if i.type.lower() == 'пластификатор':
                old_mass_films += i.mass

        for i in self.recepture.recepture_data.list_of_film_objects:
            list_new_reactives.append(i.name)
            list_new_category.append('Films')
            list_films.append(i.name)
            new_mass_film = i.mass * mass_films / (
                    old_mass_films * i.suhoi)
            list_new_mass.append(new_mass_film)
            list_add_mass_film.append(
                (old_all_mass_pigms * new_mass_film / mass_pigms) - i.mass)

        for i in self.recepture.recepture_data.list_of_additive_objects:
            if i.type.lower() == 'пластификатор':
                list_new_reactives.append(i.name)
                list_films.append(i.name)
                list_new_category.append('Additives')
                new_mass_film = i.mass * mass_films / (
                        old_mass_films * i.suhoi)
                list_new_mass.append(new_mass_film)
                list_add_mass_film.append(
                    (old_all_mass_pigms * new_mass_film / mass_pigms) - i.mass)
        # конец расчета пленок

        # print(list_films)
        # print(list_add_mass_film)
        check_list = []
        for i in list_add_mass_film:
            if int(i) > 0:
                check_list.append(True)
            else:
                check_list.append(False)

        converted_list_new_mass = list(map(normalize_number, list_new_mass))

        delete_chield(self.verticalLayout_3)
        label = QtWidgets.QLabel(self.result_w)
        label.setText("Результат")
        label.setFont(generate_font(12, bold=True))
        self.verticalLayout_3.addWidget(label, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)

        w, lo = create_w_lo(self.result_w, self.verticalLayout_3)
        label = QtWidgets.QLabel(w)
        label.setText(f'Выбранная СП: {tarir_degree_pigm}')
        lo.addWidget(label)

        for reactives, mass in zip(list_new_reactives, converted_list_new_mass):
            w, lo = create_w_lo(self.result_w, self.verticalLayout_3)
            label = QtWidgets.QLabel(w)
            label.setText(reactives)
            lo.addWidget(label)

            result_e = CustomEntry(w, padding=False)
            result_e.setMaximumSize(50, 22)
            result_e.setReadOnly(True)
            result_e.setText(mass)
            lo.addWidget(result_e)

        if all(check_list):
            w, lo = create_w_lo(self.result_w, self.verticalLayout_3)
            label = QtWidgets.QLabel(w)
            label.setText('Если есть готовый образец исходной рецептуры,\n'
                          'то вы можете на 100г рецептуры добавить:')
            lo.addWidget(label)

            converted_list_add_mass_film = list(map(normalize_number, list_add_mass_film))
            for film, mass in zip(list_films, converted_list_add_mass_film):
                w, lo = create_w_lo(self.result_w, self.verticalLayout_3)
                label = QtWidgets.QLabel(w)
                label.setText(film)
                lo.addWidget(label)

                result_e = CustomEntry(w, padding=False)
                result_e.setMaximumSize(50, 22)
                result_e.setReadOnly(True)
                result_e.setText(mass)
                lo.addWidget(result_e)

        spacerItem = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                            QtWidgets.QSizePolicy.Policy.Expanding)
        self.verticalLayout_3.addItem(spacerItem)


class CountReceptureCombo(CountReceptureConstant):
    def __init__(self, parent: ReceptureWindow):
        super(CountReceptureCombo, self).__init__(parent)
        self.setWindowTitle("Комбинированный расчет рецептуры")
        label = QtWidgets.QLabel(self.plot_w)
        label.setText("Укажите характеристики")
        label.setFont(generate_font(12, bold=True))
        self.verticalLayout_2.insertWidget(0, label, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)
        w, lo = insert_w_lo(1, self.plot_w,self.verticalLayout_2)
        label = QtWidgets.QLabel(w)
        label.setText("Содержание нелетучих веществ, %:")
        lo.addWidget(label)
        self.goal_suhoi = CustomEntry(w, padding=False)
        self.goal_suhoi.setMaximumSize(50, 22)
        self.goal_suhoi.setValidator(get_numeric_validator())
        lo.addWidget(self.goal_suhoi)
        lo.setContentsMargins(0,0,0,0)
        lo.addItem(get_h_spacer())

        w, lo = insert_w_lo(1, self.plot_w, self.verticalLayout_2)
        label = QtWidgets.QLabel(w)
        label.setText("Требуемая укрывистость мокрой пленки, г/м²:")
        lo.addWidget(label)
        self.goal_hiding = CustomEntry(w, padding=False)
        self.goal_hiding.setMaximumSize(50, 22)
        self.goal_hiding.setValidator(get_numeric_validator())
        lo.addWidget(self.goal_hiding)
        lo.addItem(get_h_spacer())
        lo.setContentsMargins(0, 0, 0, 0)


    def closeEvent(self, event):
        self.recepture.count_recepture_combo = None

    def onclick(self, event):
        tarir_degree_pigm = Decimal(float(event.xdata)).quantize(Decimal("1.00"), "ROUND_HALF_EVEN")
        # print('Мы выбрали СП на тарир ' + str(tarir_degree_pigm))
        # print('Маслоемкость тарировочной кривой ' + str(self.tarir_maslo))

        hiding_pigm = self.recepture.recepture_data.hiding_pigm
        tarir_hiding_pigm = self.tarir_hiding_pigm
        suhoi = Decimal(self.goal_suhoi.text().replace(',', '.'))
        hiding_lkm = Decimal(self.goal_hiding.text().replace(',', '.'))

        # print('Укрывистость пигменты - ' + str(hiding_pigm))
        mass_pigms = hiding_pigm * Decimal(100) / hiding_lkm

        list_new_reactives = []
        list_new_mass = []
        list_new_category = []
        list_add_mass_film = []
        list_films = []

        # начало расчета пигментов
        old_all_mass_pigms = Decimal(0)
        for i in self.recepture.recepture_data.list_of_pigment_objects:
            old_all_mass_pigms += i.mass
        for i in self.recepture.recepture_data.list_of_filler_objects:
            old_all_mass_pigms += i.mass
        for i in self.recepture.recepture_data.list_of_pigmpast_objects:
            old_all_mass_pigms += i.mass * i.suhoi_pigm

        maslo_pigm = Decimal(0)
        for i in self.recepture.recepture_data.list_of_pigment_objects:
            list_new_reactives.append(i.name)
            list_new_category.append('Pigments')
            list_new_mass.append(i.mass * mass_pigms / old_all_mass_pigms)
            maslo_pigm += i.mass * i.maslo / old_all_mass_pigms
        for i in self.recepture.recepture_data.list_of_pigmpast_objects:
            list_new_reactives.append(i.name)
            list_new_category.append('PigmPast')
            list_new_mass.append(i.mass * mass_pigms / old_all_mass_pigms)
            maslo_pigm += i.mass * i.suhoi_pigm * i.maslo / old_all_mass_pigms
        # print('Маслоемкость рассчитываемой рецептуры ' + str(maslo_pigm))
        # check_hiding_pigm = hiding_pigm < tarir_hiding_pigm
        # check_maslo_pigm = maslo_pigm < self.tarir_maslo
        # check_philum = maslo_pigm * hiding_pigm < self.tarir_maslo * self.tarir_hiding_pigm
        tarir_const_nap = tarir_degree_pigm * self.tarir_maslo
        # print('Константа наполнения по тарир ' + str(tarir_const_nap))
        # const_nap = tarir_degree_pigm * maslo_pigm
        # check_const_nap = const_nap < tarir_const_nap
        new_degree_pigm = tarir_const_nap / maslo_pigm

        # print('требуемая степень пигментирования ' + str(new_degree_pigm))
        # print('Масса пигментов ' + str(mass_pigms))
        mass_fillers = ((new_degree_pigm * suhoi) / (Decimal(1) + new_degree_pigm)) - mass_pigms
        # print('Масса напонителей ' + str(mass_fillers))
        if len(self.recepture.recepture_data.list_of_filler_objects) == 2:
            list_new_reactives.append(self.recepture.recepture_data.list_of_filler_objects[0].name)
            list_new_category.append('Fillers')
            mass_1_filler = mass_fillers * (maslo_pigm - self.recepture.recepture_data.list_of_filler_objects[1].maslo) / (
                        self.recepture.recepture_data.list_of_filler_objects[0].maslo - self.recepture.recepture_data.list_of_filler_objects[1].maslo)
            list_new_mass.append(mass_1_filler)

            list_new_reactives.append(self.recepture.recepture_data.list_of_filler_objects[1].name)
            list_new_category.append('Fillers')
            mass_2_filler = mass_fillers - mass_1_filler
            list_new_mass.append(mass_2_filler)
        elif len(self.recepture.recepture_data.list_of_filler_objects) == 1:
            list_new_reactives.append(self.recepture.recepture_data.list_of_filler_objects[0].name)
            list_new_category.append('Fillers')
            list_new_mass.append(mass_fillers)
        else:
            list_new_reactives.append('Наполнители')
            list_new_category.append('Fillers')
            list_new_mass.append(mass_fillers)
        # конец расчета пигментов и наполнителей

        mass_films = suhoi - mass_pigms - mass_fillers
        # начало расчета пленок
        for i in self.recepture.recepture_data.list_of_pigmpast_objects:  # вычитаем то что уже в пигм пастах
            mass_films -= i.mass * i.suhoi_film

        old_mass_films = Decimal(0)
        for i in self.recepture.recepture_data.list_of_film_objects:
            old_mass_films += i.mass
        for i in self.recepture.recepture_data.list_of_additive_objects:
            if i.type.lower() == 'пластификатор':
                old_mass_films += i.mass

        for i in self.recepture.recepture_data.list_of_film_objects:
            list_new_reactives.append(i.name)
            list_new_category.append('Films')
            new_mass_film = i.mass * mass_films / (
                    old_mass_films * i.suhoi)
            list_new_mass.append(new_mass_film)

        for i in self.recepture.recepture_data.list_of_additive_objects:
            if i.type.lower() == 'пластификатор':
                list_new_reactives.append(i.name)
                list_new_category.append('Additives')
                new_mass_film = i.mass * mass_films / (
                        old_mass_films * i.suhoi)
                list_new_mass.append(new_mass_film)

        # конец расчета пленок
        converted_list_new_mass = list(map(normalize_number, list_new_mass))

        delete_chield(self.verticalLayout_3)
        label = QtWidgets.QLabel(self.result_w)
        label.setText("Результат")
        label.setFont(generate_font(12, bold=True))
        self.verticalLayout_3.addWidget(label, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)

        w, lo = create_w_lo(self.result_w, self.verticalLayout_3)
        label = QtWidgets.QLabel(w)
        label.setText(f'Выбранная СП: {tarir_degree_pigm}')
        lo.addWidget(label)

        for reactives, mass in zip(list_new_reactives, converted_list_new_mass):
            w, lo = create_w_lo(self.result_w, self.verticalLayout_3)
            label = QtWidgets.QLabel(w)
            label.setText(reactives)
            lo.addWidget(label)

            result_e = CustomEntry(w, padding=False)
            result_e.setMaximumSize(50, 22)
            result_e.setReadOnly(True)
            result_e.setText(mass)
            lo.addWidget(result_e)

        self.verticalLayout_3.addItem(get_v_spacer())


class PhilumWindow(QtWidgets.QWidget):
    def __init__(self, parent: ReceptureWindow):
        super(PhilumWindow, self).__init__()
        set_window_icon(self)
        self.recepture = parent
        self.setMinimumSize(400, 400)
        self.setWindowTitle("Справочные филумы компонентов")
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.horizontalLayout.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout.setSpacing(0)
        self.setObjectName("CountRecepture")
        self.setStyleSheet("""
                QWidget#CountRecepture{
                background: #f9f9f9;
                }
                """)
        self.philum_area = QtWidgets.QTextEdit(parent=self)
        self.horizontalLayout.addWidget(self.philum_area)
        self.philum_area.setReadOnly(True)
        self.list_name = ['1.Красный и чёрный железоокисные ', '2.Окись хрома пигментная ',
                          '3.Сурик железный (высший сорт)', '4.Свинцовая зелень', '5.Титанат хрома дегидратированный',
                          '6.Жёлтый железо-цинковый ', '7.Кадмиевый жёлтый (импорт)',
                          '8.Коричневый на основе оксида кобальта', '9.Сажа термическая', '10.Крон свинцовый оранжевый',
                          '11.Крон свинцовый жёлтый(лимонный)', '12.Коричневый железоокисный ',
                          '13.Охра красная натуральная', '14.Железная лазурь (милори)', '15.Ртутно-кадмиевый красный',
                          '16.Титанат кобальта зелёный', '17.Сажа ламповая', '18.Кадмиевый красный (импорт)',
                          '19.Диоксид титана рутил (импорт)', '20.Жёлтый железоокисный',
                          '21.Умбра натуральная тёмно-коричневая', '22.Мумия природная железоокисная',
                          '23.Крон свинцовый молибдатный', '24.Киноварь', '25.Кобальт зелёный',
                          '26.Органический красный С', '27.Фталоцианин. голубой 2 «З»У / ПАП-1', '28.Кадмий оранжевый',
                          '29.Марганцевая коричневая', '30.Зелёный органический', '31.Сажа канальная',
                          '32.Венецианская и индийск. красные жел.ок', '33.Сиена жжёная', '34.Бордо СМ',
                          '35.Фталоцианиновый зелёный', '36.Феодосийская коричневая',
                          '37.Алый лакокрасочный имп./отеч.', '38.Титано-кобальто-алюминиевый синий',
                          '39.Титанат никеля жёлтый', '40.Марганцевый фиолетовый', '41.Синий антрахиноновый ОА',
                          '42.Лак красный 2СМ', '43.Диоксид титана анатаз', '44.Хром-кобальт зелёно-голубой',
                          '45.Силикохромат свинца', '46.Умбра жжёная', '47.Сажа ацетиленовая',
                          '48.Фиолетовый хинакридоновый', '49.Ярко-оранжевый антрахиноновый К',
                          '50.Мумия бокситная жёлтая', '51.Церулеум (небесно-голубой)',
                          '52.Охра жёлтая тёмная натуральная', '53.Кобальт фиолетовый тёмный', '54.Бордо К',
                          '55.Хромат стронция', '56.Ярко-оранжевый антрахиноновый', '57.Марганцовая голубая',
                          '58.Хромат бария', '59.Литопон (ZnS + BaSO4)', '60.Цинковые белила',
                          '61.Фталоцианин. голубой / красный 5с', '62.Охра красная прокалённая',
                          '63.Кассельская прокалённая', '64.Охра коричневая', '65.Карбонатные свинцовые белила',
                          '66.Умбра натуральная светлая', '67.Марс коричневый тёмный', '68.Охра жёлтая средняя',
                          '69.Охра жёлтая золотистая', '70.Марс коричневый светлый', '71.Охра жёлтая светлая',
                          '72.Мумия глинистая красная', '73.Архангельская коричневая', '74.Хромат бария-калия',
                          '75.Титанат хрома гидратированный', '76.Органический жёлтый 2 «З»У',
                          '77.Органический жёлтый светопрочный', '78.Основной хромат цинка-калия',
                          '79.Оранжевый органическ. с наполнителем', '80.Тетраоксихромат свинца',
                          '81.Марс жёлтый прозрачный', '82.Волконскоитовая зелёная', '83.Изумрудная зелень',
                          '84.Цинковый крон', '85.Минеральная коричневая светлая', '86.Минеральная коричневая тёмная',
                          '87.Кобальт синий', '88.Стронциановая жёлтая', '89.Синий силикат кобальта-цинка',
                          '90.Ультрамарин фиолетовый', '91.Ультрамарин синий', '92.Кобальт фиолетовый светлый',
                          '93.Мел (наполнитель)']
        self.list_value = ['1,2-2,8', '1,5-2,8', '1,5-5,0', '1,6-4,8', '1,9-3,5', '2,4-5,0', '2,4-8,1', '2,5-4,2',
                           '2,5-4,4', '2,7-5,4', '2,7-15,0', '3', '3,0-5,0', '3,0-7,2', '3,0-7,5', '3,6-7,5', '4-4,8',
                           '4-5,5', '4-7', '4,2-14,0', '4,4-22,5', '4,5-9,1', '4.8', '5', '5-12', '5,5-7,7', '5,6/5,6',
                           '5,6-16,2', '5,6-18', '6', '6', '6,0-9,0', '6,0-14,0', '6.4', '6.6', '6.6', '6,6-8,0/9,5',
                           '6,6-13,2', '6,84-8,0', '6,8-13,5', '7', '7', '7-10', '7,0-11,25', '7,15-9,0', '7,2-13,8',
                           '7.5', '7.68', '8.4', '8,4-10,8', '9,0-14,0', '9,0-18,0', '9,75-26,3', '10.5', '10,5-22,5',
                           '11', '11,2-14,4', '12-21,6', '12,1-21,0', '13,2-22,4', '13,5 / 13,6', '13,8-16,3', '14-28',
                           '14-30', '14,4-24,0', '15-20', '15-26,4', '15-28', '16,0-32,5', '16,3-28', '16,3-36',
                           '16,5-29,6', '17,5-35', '18-27', '18,2-35,1', '20', '20.4', '20-32,5', '20.8', '20-42',
                           '20-90', '22,5-37,5', '27,8-31,5', '24-126', '26-32', '30-36', '30-54', '30,1-42,3', '35-56',
                           '35-55', '35-60', '96-170', '120']
        philum_str = ""
        for name, value in zip(self.list_name, self.list_value):
            philum_str+= f"{name}: {value}\n"

        self.philum_area.setText(philum_str)

    def closeEvent(self, event):
        self.recepture.philum_window = None


class SaveAsWindow(QtWidgets.QWidget):
    def __init__(self, parent: ReceptureWindow, project, iteration, save_callback):
        super(SaveAsWindow, self).__init__()
        set_window_icon(self)
        self.project = project
        self.iteration = iteration
        self.parent_obj = parent
        self.save_callback = save_callback
        self.resize(300, 96)
        self.setMaximumSize(QtCore.QSize(300, 150))
        self.gridLayout = QtWidgets.QGridLayout(self)
        self.gridLayout.setObjectName("gridLayout")
        self.iterl = QtWidgets.QLabel(parent=self)
        self.iterl.setText("Выберите итерацию:")
        self.gridLayout.addWidget(self.iterl, 0, 0, 1, 1)
        self.iter_c = CustomCombobox(parent=self)
        self.iter_c.addItems(self.get_list_iter())
        self.iter_c.setText(self.iteration)
        self.iter_c.setObjectName("iter_c")
        self.gridLayout.addWidget(self.iter_c, 0, 1, 1, 1)
        self.name_l = QtWidgets.QLabel(parent=self)
        self.name_l.setText("Новое название:")
        self.gridLayout.addWidget(self.name_l, 1, 0, 1, 1)
        self.name_e = CustomEntry(parent=self, padding=False)
        self.gridLayout.addWidget(self.name_e, 1, 1, 1, 1)

        self.save_b = ColorButton(parent=self, color="blue")
        self.save_b.clicked.connect(lambda : self.save())
        self.gridLayout.addWidget(self.save_b, 2, 0, 1, 2)
        self.save_b.setText("Сохранить")
        self.setWindowTitle("Сохранить как")

    def get_list_iter(self) -> List[str]:
        iters = os.listdir('saves/' + self.project)
        iters.remove('params')
        return iters

    def get_list_recepture(self, iter) -> List[str]:
        with SqliteDict('saves/' + self.project + '/' + iter) as mydict:
            return list(mydict.keys())

    def closeEvent(self, event):
        self.parent_obj.save_as_window = None

    def save(self):
        iter = self.iter_c.text().strip()
        new_name = self.name_e.text().strip()
        if iter == "" or new_name == "":
            InfoWindow("Имя не указано").exec()
            return
        list_names = self.get_list_recepture(iter)
        if new_name in list_names:
            InfoWindow("Рецептура с таким именем уже существует.").exec()
            return

        self.save_callback(iter, new_name)
        self.closeEvent(event=None)
        self.destroy()


class WordExport:
    def __init__(self, data_model: ReceptureDataModel):
            self.data_model = data_model
            self.db = DB()
            self.name = self.data_model.name.strip()
            self.component_list = self.get_components_list()
            self.experiment_list = self.data_model.experiment_list
            self.note = self.data_model.notes
            self.mass = self.data_model.mass
            count_dict = self.data_model.get_count_dict()
            count_names = list(count_dict.keys())
            count_names.sort()




            dialog = QFileDialog()
            dialog.setWindowTitle("Выберите папку для сохранения")
            dialog.setFileMode(QFileDialog.FileMode.Directory)
            dialog.setOptions(QFileDialog.Option.DontUseNativeDialog)
            if dialog.exec():
                new_file = dialog.selectedFiles()[0]
                new_file = f"{new_file}\\{self.name}.docx"
            else:
                return

            forma = 'files\\forma.docx'
            shutil.copy2(forma, new_file)
            # new_file = os.path.abspath(f'documents\{self.name}.docx')

            word = client.gencache.EnsureDispatch('Word.Application')
            doc = word.Documents.Open(new_file)
            table = doc.Tables(2)  # создаем строки в  таблице компонентов
            for i in range(len(self.component_list) + 1):
                table.Rows.Add()

            table = doc.Tables(3)  # создаем строки в экспериментальной таблице
            for i in range(len(self.experiment_list)):
                table.Rows.Add()

            table = doc.Tables(4)  # создаем строки в расчетной таблице
            for i in range(len(count_names)):
                table.Rows.Add()

            doc.Close(True)
            word.Application.Quit()

            doc = docx.Document(new_file)

            doc.paragraphs[0].add_run(self.name)
            tables = doc.tables

            # заполняем таблицу данными
            for row in range(len(self.component_list)):
                # получаем ячейку таблицы
                cell = tables[1].cell(row + 1, 0)
                # записываем в ячейку данные
                cell.text = self.component_list[row][0]

                cell = tables[1].cell(row + 1, 1)
                # записываем в ячейку данные
                cell.text = self.component_list[row][1]

            cell = tables[1].cell(row + 2, 0)
            cell.text = "Итого:"
            cell = tables[1].cell(row + 2, 1)
            cell.text = str(self.mass).replace('.', ',')

            for row in range(len(self.experiment_list)):
                # получаем ячейку таблицы
                cell = tables[2].cell(row + 1, 0)
                # записываем в ячейку данные
                cell.text = self.experiment_list[row][0]

                cell = tables[2].cell(row + 1, 2)
                # записываем в ячейку данные
                cell.text = self.experiment_list[row][1]

                cell = tables[2].cell(row + 1, 1)
                # записываем в ячейку данные
                cell.text =  self.experiment_list[row][2]

            for row in range(len(count_names)):
                # получаем ячейку таблицы
                cell = tables[3].cell(row + 1, 0)
                # записываем в ячейку данные
                cell.text = count_names[row]
                value = count_dict[count_names[row]]
                if isinstance(value, Decimal):
                    value = normalize_number(value)
                if isinstance(value, float):
                    value = str(round(value, 2)).replace(".", ",")

                if count_names[row] == "Степень пигментирования":
                    cell = tables[3].cell(row + 1, 1)
                    # записываем в ячейку данные
                    cell.text = value + " : 1"
                else:
                    cell = tables[3].cell(row + 1, 1)
                    # записываем в ячейку данные
                    cell.text = value

            cell = tables[0].cell(0, 0)
            note = '\n'.join(wrap(self.note, 25))
            cell.text = 'Заметки:\n' + note

            doc.save(new_file)
            os.startfile(new_file)

    def get_components_list(self):
        components_list = self.data_model.component_list
        category_list = self.data_model.category_list
        if self.data_model.flag_2k:
            components_list += self.data_model.component_list_2
            category_list += self.data_model.category_list_2

        normalize_components_list = []
        for component in components_list:
            if isinstance(component, str):
                comment = component
                normalize_components_list.append((comment, ""))
            else:
                normalize_components_list.append(component)
        components_list = normalize_components_list


        if InfoWindow("Заменить названия компонентов \nна шифр?").exec():
            coded_components_list = []
            for component, category in zip(components_list, category_list):

                name = component[0]
                mass = component[1]
                if category.strip() != "":
                    code = self.db.get_info_reactive(category, name, 'code')
                    if len(code) > 0:
                        code = code[0][0]
                    else:
                        code = name
                else:
                    code = name
                coded_components_list.append((code, mass))

            components_list = coded_components_list

        return components_list



    def save(self):
        pass
